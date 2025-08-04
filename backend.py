import os
import mimetypes
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
os.environ["OMP_NUM_THREADS"] = "1"
os.environ["PATH"] += os.pathsep + r"C:\ffmpeg\bin\bin"
os.environ["FFMPEG_BINARY"] = r"C:\ffmpeg\bin\bin\ffmpeg.exe"
import re
import json
from datetime import datetime, timedelta
from flask import Flask, request, jsonify, send_from_directory, g, session
from flask_cors import CORS
from flask_mail import Mail, Message
import pyotp
import secrets
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from prompts import (
    CORE_INSTRUCTION,
    FORMAT_PROMPT,
    DETAILED_RESPONSE_INSTRUCTION,
    FILTER_PROMPT,
    GREETING_INSTRUCTION,
    SUMMARIZATION_INSTRUCTION,
    COMPARISON_INSTRUCTION,
    EXTRACTION_INSTRUCTION,
    ANALYSIS_INSTRUCTION,
    TECHNICAL_INSTRUCTION,
    DEFINITION_INSTRUCTION,
    EXAMPLE_INSTRUCTION,
    TRANSLATION_INSTRUCTION,
    CLARIFICATION_INSTRUCTION,
    GENERAL_INSTRUCTION,
    DPDP_INSTRUCTION,
    PARLIAMENT_INSTRUCTION,
    HOW_TO_USE,
    EXAMPLE_USAGE,
    generate_source_link
)
from backend.utils import initialize_faiss_index, update_faiss_index
from backend.banner_utils import save_banner_file, load_banner_file, remove_banner_file, get_banner_storage_info
import logging
import faiss
import warnings
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import glob
import uuid
from pypdf import PdfReader
from werkzeug.utils import secure_filename
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from elevenlabs import play # Keep play for local playback if needed
from elevenlabs.client import ElevenLabs # Import the client class
import requests
from ollama import Client as OllamaClient
import threading
import pandas as pd
import pdfplumber
import pytesseract
import base64
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
from PIL import Image
import cv2
import numpy as np
import json as pyjson
# PATCH: Add Whisper for server-side audio transcription
try:
    import whisper
    whisper_model = whisper.load_model('base')
    WHISPER_AVAILABLE = True
except Exception as e:
    WHISPER_AVAILABLE = False
    whisper_model = None
    print(f"[WARNING] Whisper not available: {e}")

# Suppress Faiss GPU warnings
warnings.filterwarnings("ignore", message=".*GpuIndexIVFFlat.*")

# Configure logging
logging.basicConfig(level=logging.INFO) # Changed from DEBUG to INFO
logger = logging.getLogger(__name__)

# Suppress specific loggers
logging.getLogger('werkzeug').setLevel(logging.WARNING)  # Reduce werkzeug logging
logging.getLogger('watchdog').setLevel(logging.WARNING)  # Reduce watchdog logging
logging.getLogger('faiss').setLevel(logging.WARNING)  # Reduce faiss logging
logging.getLogger('langchain').setLevel(logging.WARNING)  # Reduce langchain logging

# Explicitly configure Faiss to use CPU
faiss.omp_set_num_threads(4)  # Set number of threads for CPU operations

# Initialize Flask app
app = Flask(__name__)

# Load environment variables FIRST
env_path = os.path.join(os.path.dirname(__file__), '.env')
print(f"Looking for .env file at: {env_path}")
print(f"File exists: {os.path.exists(env_path)}")
load_dotenv(env_path)

# Debug: Check what environment variables are loaded
print(f"EMAIL_USERNAME from env: {os.getenv('EMAIL_USERNAME')}")
print(f"EMAIL_PASSWORD from env: {os.getenv('EMAIL_PASSWORD', 'NOT_FOUND')}")

# Configure Flask app
app.config['DEBUG'] = False
app.config['TEMPLATES_AUTO_RELOAD'] = False

# Email configuration for OTP
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.getenv('EMAIL_USERNAME', 'your-email@gmail.com')
app.config['MAIL_PASSWORD'] = os.getenv('EMAIL_PASSWORD', 'your-app-password')
app.config['MAIL_DEFAULT_SENDER'] = os.getenv('EMAIL_USERNAME', 'your-email@gmail.com')

# Debug: Print email configuration (remove in production)
logger.info(f"Email configuration loaded:")
logger.info(f"MAIL_USERNAME: {app.config['MAIL_USERNAME']}")
logger.info(f"MAIL_PASSWORD: {'*' * len(app.config['MAIL_PASSWORD']) if app.config['MAIL_PASSWORD'] != 'your-app-password' else 'NOT_SET'}")
logger.info(f"MAIL_DEFAULT_SENDER: {app.config['MAIL_DEFAULT_SENDER']}")

# Initialize Flask-Mail
mail = Mail(app)

# OTP storage (in production, use Redis or database)
otp_storage = {}
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'supersecretkey')
app.config['SESSION_COOKIE_NAME'] = 'rag_chatbot_session_id'

# Configure CORS
CORS(app, resources={
    r"/*": {
        "origins": ["http://localhost:3000"],
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Accept", "Authorization"],
        "expose_headers": ["X-PDF-Page", "X-PDF-Section"],
        "supports_credentials": True,
        "max_age": 3600
    }
})

# Create uploads directory if it doesn't exist
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'backend', 'pdfs')
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
logger.info(f"UPLOAD_FOLDER resolved to: {UPLOAD_FOLDER}")

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# Environment variables already loaded above

# Configure Gemini API
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=GOOGLE_API_KEY)

# Configure ElevenLabs API
ELEVENLABS_API_KEY = os.getenv("ELEVENLABS_API_KEY")
# Initialize ElevenLabs client
elevenlabs_client = ElevenLabs(api_key=ELEVENLABS_API_KEY)

# Configure Sarvam API
SARVAM_API_KEY = os.getenv("SARVAM_API_KEY")

# Initialize the LLM
llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash-preview-05-20",  # Updated to correct model name
    temperature=0.7,
    google_api_key=GOOGLE_API_KEY,
    streaming=True
)

# Initialize the embeddings
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")  # Multilingual support

document_heading = None

# Path to the manual page mappings file (commented out for auto-detection test)
# PDF_PAGE_MAPPINGS_FILE = os.path.join(os.path.dirname(__file__), 'pdf_page_mappings.json')

# def load_manual_page_mappings():
#     """Loads manual page number mappings from a JSON file."""
#     if os.path.exists(PDF_PAGE_MAPPINGS_FILE):
#         try:
#             with open(PDF_PAGE_MAPPINGS_FILE, 'r') as f:
#                 return json.load(f)
#         except Exception as e:
#             logger.error(f"Error loading manual page mappings from {PDF_PAGE_MAPPINGS_FILE}: {e}")
#     return {}

# manual_page_mappings = load_manual_page_mappings()

def load_pdfs_from_directory(directory_path: str) -> list[Document]:
    """Load all PDFs from a directory and convert them to documents"""
    global document_heading
    documents = []
    
    # Add logging to debug PDF loading
    logger.info(f"load_pdfs_from_directory called with directory_path: {directory_path}") # Log the input path
    pdf_pattern = os.path.join(app.config['UPLOAD_FOLDER'], "*.pdf") # Use configured UPLOAD_FOLDER
    logger.info(f"Searching for PDFs with pattern: {pdf_pattern}")
    pdf_files = glob.glob(pdf_pattern)
    logger.info(f"Found PDF files: {pdf_files}") # Log the result of glob.glob
    
    for pdf_file in pdf_files:
        try:
            reader = PdfReader(pdf_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            doc = Document(
                page_content=text,
                metadata={"source": os.path.basename(pdf_file)}
            )
            documents.append(doc)
            logger.info(f"Successfully loaded PDF: {pdf_file}")
        except Exception as e:
            logger.error(f"Error loading PDF {pdf_file}: {str(e)}")
    
    if documents:
        # Set the global document_heading using the first loaded document's source metadata
        document_heading = documents[0].metadata.get('source', 'Uploaded Document')
        logger.info(f"Document Heading set to: {document_heading}")
    else:
        document_heading = 'No Document Loaded'
        logger.warning("No PDF documents loaded from directory.")
    
    return documents

def process_pdf(filepath):
    """Process a single PDF file and update the FAISS index"""
    global retriever
    try:
        # Read the PDF file
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        # Create a document with metadata
        doc = Document(
            page_content=text,
            metadata={"source": os.path.basename(filepath)}
        )
        
        # Initialize or update the FAISS index
        if retriever is None:
            # If no retriever exists, create a new FAISS index
            vectorstore = initialize_faiss_index([doc], embeddings)
            retriever = vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 20}
            )
        else:
            # If retriever exists, update the existing FAISS index
            update_faiss_index([doc], embeddings)
            
        logger.info(f"Successfully processed PDF: {filepath}")
        return True
    except Exception as e:
        logger.error(f"Error processing PDF {filepath}: {str(e)}")
        return False

# Initialize vector store with PDFs
pdf_documents = load_pdfs_from_directory("backend/pdfs")
vectorstore = initialize_faiss_index(pdf_documents, embeddings)
retriever = vectorstore.as_retriever(
    search_type="similarity",
    search_kwargs={"k": 20}
)

# In-memory session storage
sessions = {}

# Session file path
SESSIONS_FILE = 'sessions.json'

MAX_CONVERSATION_HISTORY = int(os.getenv('MAX_CONVERSATION_HISTORY', 10))

def save_sessions():
    """Save sessions to a JSON file."""
    try:
        with open(SESSIONS_FILE, 'w') as f:
            json.dump(sessions, f, indent=4)
        logger.info("Sessions saved to file.")
    except Exception as e:
        logger.error(f"Error saving sessions: {e}")

def load_sessions():
    """Load sessions from a JSON file."""
    global sessions
    if os.path.exists(SESSIONS_FILE):
        try:
            with open(SESSIONS_FILE, 'r') as f:
                sessions = json.load(f)
            logger.info("Sessions loaded from file.")
        except Exception as e:
            logger.error(f"Error loading sessions: {e}")

@app.route('/sessions', methods=['GET'])
def get_sessions():
    """Get all sessions"""
    return jsonify(list(sessions.values()))

@app.route('/sessions', methods=['POST'])
def create_session():
    """Create a new session"""
    session_id = str(uuid.uuid4())
    session = {
        'id': session_id,
        'title': f'New Chat {len(sessions) + 1}',
        'createdAt': datetime.now().isoformat(),
        'messages': [],
        'agentId': None,  # Initialize agentId as None for new sessions
        'agentName': None  # Initialize agentName as None for new sessions
    }
    sessions[session_id] = session
    save_sessions() # Save sessions after creating a new one
    
    # Log monitoring data
    try:
        data = request.json or {}
        user_email = data.get('userEmail', 'anonymous@user.com')
        agent_id = data.get('agentId')
        
        # Update session with agent information if provided
        if agent_id:
            sessions[session_id]['agentId'] = agent_id
            if agent_id in AGENTS_DATA:
                sessions[session_id]['agentName'] = AGENTS_DATA[agent_id].get('name', 'Unknown Agent')
            else:
                sessions[session_id]['agentName'] = 'Unknown Agent'
            save_sessions()  # Save the updated session
        
        log_user_activity(
            user_email=user_email,
            action='create_session',
            session_id=session_id,
            agent_id=agent_id
        )
    except Exception as e:
        logger.error(f"Error logging session creation: {e}")
    
    return jsonify(session)

@app.route('/sessions/<session_id>', methods=['DELETE', 'OPTIONS'])
def delete_session(session_id):
    """Delete a session"""
    if request.method == 'OPTIONS':
        return '', 200
        
    if session_id in sessions:
        # Log monitoring data before deletion
        try:
            data = request.json or {}
            user_email = data.get('userEmail', 'anonymous@user.com')
            session_data = sessions[session_id]
            
            # Log session metrics
            message_count = len(session_data.get('messages', []))
            created_at = datetime.fromisoformat(session_data.get('createdAt', datetime.now().isoformat()))
            session_duration = (datetime.now() - created_at).total_seconds()
            
            log_session_metrics(
                session_id=session_id,
                user_email=user_email,
                agent_id=session_data.get('agentId'),
                message_count=message_count,
                session_duration=session_duration,
                model_used='unknown'  # Could be extracted from messages if needed
            )
            
            log_user_activity(
                user_email=user_email,
                action='delete_session',
                session_id=session_id,
                agent_id=session_data.get('agentId')
            )
        except Exception as e:
            logger.error(f"Error logging session deletion: {e}")
        
        del sessions[session_id]
        save_sessions() # Save sessions after deleting one
        return '', 204
    return jsonify({'error': 'Session not found'}), 404

@app.route('/sessions/<session_id>', methods=['PUT'])
def update_session(session_id):
    """Update a session's information"""
    if session_id not in sessions:
        return jsonify({'error': 'Session not found'}), 404
    
    data = request.json
    new_title = data.get('title')
    
    if new_title is not None:
        sessions[session_id]['title'] = new_title
        save_sessions() # Save sessions after updating one
        return jsonify(sessions[session_id])
    
    return jsonify({'error': 'No update data provided'}), 400

@app.route('/sessions/<session_id>/messages', methods=['GET'])
def get_messages(session_id):
    """Get messages for a session"""
    if session_id not in sessions:
        return jsonify({'error': 'Session not found'}), 404
    return jsonify(sessions[session_id]['messages'])

def get_label_to_page_map(pdf_path):
    """Extract a mapping from page label (logical page number/section) to actual page number."""
    try:
        reader = PdfReader(pdf_path)
        label_to_page = {}
        page_labels = getattr(reader, 'page_labels', None)
        if page_labels is None and hasattr(reader, 'get_page_labels'):
            page_labels = reader.get_page_labels()
        
        if page_labels is None:
            print("DEBUG: [get_label_to_page_map] No page labels found with pypdf, falling back to 1-based page numbers.")
            for i in range(len(reader.pages)):
                label_to_page[str(i + 1)] = i + 1
        else:
            print(f"DEBUG: [get_label_to_page_map] Found pypdf page labels: {page_labels}")
            for i, label in enumerate(page_labels):
                label_to_page[str(label)] = i + 1  # 1-based page number
        print(f"DEBUG: [get_label_to_page_map] Generated pypdf label_to_page map: {label_to_page}")
        return label_to_page
    except Exception as e:
        print(f"DEBUG: [get_label_to_page_map] Error for {pdf_path}: {e}")
        logger.error(f"Error in get_label_to_page_map for {pdf_path}: {e}")
        return {}

def format_response_with_sources(response, sources):
    """Format the response with source information, always outputting PDF links in markdown (parentheses) format for frontend parsing. Tries to match section/page to the correct filename from docs. Adds &highlight=... for frontend highlighting."""
    if not response:
        return None
        
    # Split response into content and sources
    parts = response.split('**Sources:**')
    if len(parts) != 2:
        return response
        
    content = parts[0].strip()
    sources_block = parts[1].strip()

    formatted_sources = []
    # Regex to match pdf:// links
    pdf_link_regex = r'pdf://([^/]+)(?:/page/(\d+))?(?:#section=([^\s)]+))?'
    # Regex to match Gemini's [label] (Section ..., Page ...)
    gemini_source_regex = r'\[(.*?)\]\s*\(Section ([^,]+), Page (\d+)\)'
    # Regex to match [Source: ...; pdf://...]
    llama_source_regex = r'pdf://([^/]+)(?:/page/(\d+))?(?:#section=([^\s\]]+))?'

    # Use the exact PDF text as the highlight, or fallback if empty
    def get_highlight_texts(filename, page):
        filename = str(filename).strip().lower()
        page = str(page).strip() if page is not None else None
        highlights = []
        for doc in sources or []:
            doc_filename = str(doc.metadata.get('source', '')).strip().lower()
            doc_page = str(doc.metadata.get('page', '')).strip()
            if (not page or doc_page == page) and doc_filename == filename:
                text = doc.page_content.strip()
                # Split into sentences and lines
                # First split by newlines, then by sentence boundaries
                lines = [l.strip() for l in text.split('\n') if l.strip()]
                for line in lines:
                    # Split by sentence boundaries (., !, ?)
                    sentences = re.split(r'(?<=[.!?]) +', line)
                    for s in sentences:
                        s = s.strip()
                        if s:
                            highlights.append(s)
        if not highlights and sources and len(sources) > 0:
            text = sources[0].page_content.strip()
            lines = [l.strip() for l in text.split('\n') if l.strip()]
            for line in lines:
                sentences = re.split(r'(?<=[.!?]) +', line)
                for s in sentences:
                    s = s.strip()
                    if s:
                        highlights.append(s)
        print(f"[DEBUG] All highlights for {filename} page {page}: {highlights}")
        return highlights

    for line in sources_block.split('\n'):
        line = line.strip()
        if not line:
            continue
        # Try pdf://... link
        match = re.search(pdf_link_regex, line)
        if match:
            filename, page, section = match.groups()
            corrected_filename = filename
            if sources:
                for doc in sources:
                    doc_filename = doc.metadata.get('source', filename)
                    doc_page = str(doc.metadata.get('page', ''))
                    doc_section = str(doc.metadata.get('section', '')).replace(' ', '_')
                    if page and doc_page == page:
                        corrected_filename = doc_filename
                        break
            link = f"- (pdf://{corrected_filename}"
            if page:
                link += f"/page/{page}"
            if section:
                link += f"#section={section}"
            # Add highlight param using all source chunks
            highlights = get_highlight_texts(corrected_filename, page)
            if highlights:
                from urllib.parse import quote
                # Join all highlights with a delimiter
                joined = '|||'.join(highlights)
                link += f"&highlight={quote(joined)}"
            link += ")"
            formatted_sources.append(link)
            continue
        # Try Gemini format: [label] (Section ..., Page ...)
        match = re.search(gemini_source_regex, line)
        if match:
            label, section, page = match.groups()
            # Try to find doc with matching page and section
            corrected_filename = None
            highlight_text = None
            for doc in sources:
                doc_page = str(doc.metadata.get('page', ''))
                doc_section = str(doc.metadata.get('section', '')).replace(' ', '_')
                if doc_page == page and (not section or section.replace(' ', '_') in doc_section):
                    corrected_filename = doc.metadata.get('source')
                    # Use first 200 characters of the chunk as highlight
                    highlight_text = doc.page_content[:200].replace('\n', ' ')
                    break
            if not corrected_filename and sources:
                corrected_filename = sources[0].metadata.get('source', 'Unknown')
                highlight_text = sources[0].page_content[:200].replace('\n', ' ')
            link = f"- (pdf://{corrected_filename}/page/{page}#section={section.replace(' ', '_')}"
            # Add highlight param using the chunk excerpt
            if highlight_text:
                from urllib.parse import quote
                link += f"&highlight={quote(highlight_text)}"
            link += ")"
            formatted_sources.append(link)
            continue
        # Try Llama format: [Source: ...; pdf://...]
        match = re.search(llama_source_regex, line)
        if match:
            filename, page, section = match.groups()
            corrected_filename = filename
            if sources:
                for doc in sources:
                    doc_filename = doc.metadata.get('source', filename)
                    doc_page = str(doc.metadata.get('page', ''))
                    doc_section = str(doc.metadata.get('section', '')).replace(' ', '_')
                    if page and doc_page == page:
                        corrected_filename = doc_filename
                        break
            link = f"- (pdf://{corrected_filename}"
            if page:
                link += f"/page/{page}"
            if section:
                link += f"#section={section}"
            # Add highlight param using the answer text
            highlight = get_highlight_texts(corrected_filename, page)
            if highlight:
                from urllib.parse import quote
                link += f"&highlight={quote(highlight)}"
            link += ")"
            formatted_sources.append(link)
            continue
        # Fallback: append the original line
            formatted_sources.append(line)

    # Reconstruct the response with formatted sources
    formatted_response = f"{content}\n\n**Sources:**\n" + "\n".join(formatted_sources)
    print(f"[DEBUG] Final formatted response: {formatted_response[:500]}")
    return formatted_response

def create_structured_prompt(user_message, context, doc_id=None, DOCUMENT_HEADING=None, USER_NAME=None, AGENT_ID=None):
    """Create a structured prompt using the instruction constants"""
    # Check if the message contains Hindi characters
    hindi_pattern = re.compile(r'[\u0900-\u097F]')
    is_hindi = bool(hindi_pattern.search(user_message))
    
    # Check if it's a pure greeting vs a question with greeting
    greeting_keywords = ['hello', 'hi', 'hey', 'good morning', 'good afternoon', 'good evening']
    question_keywords = ['what', 'why', 'how', 'when', 'where', 'who', 'which', 'can', 'could', 'would', 'should', 'is', 'are', 'do', 'does', 'did']
    
    # Check if message contains both greeting and question keywords
    has_greeting = any(keyword in user_message.lower() for keyword in greeting_keywords)
    has_question = any(keyword in user_message.lower() for keyword in question_keywords)
    
    # Select instruction based on agent ID first
    if AGENT_ID == 'DPDP':
        selected_instruction = DPDP_INSTRUCTION
    elif AGENT_ID == 'Parliament':
        selected_instruction = PARLIAMENT_INSTRUCTION
    elif has_greeting and not has_question and len(user_message.split()) <= 3:
        # Only treat as pure greeting if it's just a greeting without any question content
        return {
            "prompt": GREETING_INSTRUCTION,
            "query_type": "greeting",
            "metadata": {"is_greeting": True, "agent_id": AGENT_ID}
        }
    else:
        # Determine if the user is asking for a comparison to prioritize table format
        is_comparison_query = any(keyword in user_message.lower() for keyword in ['compare', 'contrast', 'difference', 'similarity', 'vs'])
        
        # Determine if the user is asking for a detailed/elaborated response
        is_detailed_query = any(keyword in user_message.lower() for keyword in ['elaborate', 'detail', 'detailed', 'explain more', 'tell me more', 'in depth'])
        
        # Determine if the user is asking for a definition
        is_definition_query = any(keyword in user_message.lower() for keyword in ['define', 'what is', 'meaning of', 'definition of', 'explain'])

        # Select the appropriate instruction based on query type
        if is_detailed_query:
            selected_instruction = DETAILED_RESPONSE_INSTRUCTION
        elif is_comparison_query:
            selected_instruction = COMPARISON_INSTRUCTION
        elif is_definition_query:
            selected_instruction = DEFINITION_INSTRUCTION
        else:
            selected_instruction = CORE_INSTRUCTION

    # Create the base prompt with core instructions and formatting guidelines
    language_instruction = """
IMPORTANT: Respond in the same language as the user's question. If the question is in Hindi, respond in Hindi. If the question is in English, respond in English.
""" if is_hindi else ""

    base_prompt = f"""
{CORE_INSTRUCTION}

{language_instruction}

You are an AI assistant designed to answer questions based on the provided context.
Carefully read the following context and answer the user's question directly and concisely.
If the answer is not found in the context, state that you cannot answer based on the provided information.

## Context Source(s): {DOCUMENT_HEADING}
## Context:
{context}

## User's Question:
{user_message}

## Formatting Instructions (Strict Adherence Required):

1. **Direct Output**: Respond with the requested information or format immediately after a brief acknowledgment. Avoid conversational filler or lengthy introductions, especially when a specific format like a table is requested.

2. **Comparison Tables**: **WHEN ASKED TO COMPARE**, your response MUST start with a markdown table summarizing the comparison. Generate the table first, followed by any necessary detailed explanation. Crucially, use standard markdown table syntax ONLY. Do NOT use HTML tags like `<br>` within table cells. Ensure the complete text for each item is included in its respective cell without truncation.
   - Use markdown table syntax with proper alignment
   - Include headers with bold text
   - Example:
     | **Header 1** | **Header 2** |
     |--------------|--------------|
     | Content 1    | Content 2    |

3. **PDF Links Format**:
   - Use the format: pdf://<filename>/page/<page_number>#section=<section_name>
   - Example: pdf://DPDP_act.pdf/page/9#section=Section_1
   - The backend will automatically convert these to http:// URLs with query parameters
   - Always include source information in the response

4. **Context-Aware Responses**:
   - If the message contains both a greeting and a question, prioritize answering the question
   - Maintain professional tone while being helpful
   - Focus on the substantive content of the query

5. **GDPR and Privacy Queries**:
   - For privacy-related questions, provide detailed and accurate information
   - Include relevant legal context and implications
   - Cite specific sections of relevant documents
"""

    # Combine the selected instruction with the base prompt
    final_prompt = f"""{selected_instruction}

{base_prompt}"""

    # Determine query type based on selected instruction or other factors
    query_type = "general"
    if AGENT_ID == 'DPDP':
        query_type = "dpdp_agent"
    elif AGENT_ID == 'Parliament':
        query_type = "parliament_agent"
    elif is_detailed_query:
        query_type = "detailed"
    elif is_comparison_query:
        query_type = "comparison"

    # Return the appropriate prompt based on query type
    return {
        "prompt": final_prompt,
        "query_type": query_type,
        "metadata": {
            "document_id": doc_id,
            "document_heading": DOCUMENT_HEADING,
            "user_name": USER_NAME,
            "agent_id": AGENT_ID
        }
    }

# Define the path to agents.json
AGENTS_JSON_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'react-frontend', 'src', 'data', 'agents.json')

# Global variable to store agent data
AGENTS_DATA = {}

def save_agents_to_json():
    """Saves the current AGENTS_DATA to the agents.json file."""
    try:
        agents_list = list(AGENTS_DATA.values())
        os.makedirs(os.path.dirname(AGENTS_JSON_PATH), exist_ok=True)
        with open(AGENTS_JSON_PATH, 'w', encoding='utf-8') as f:
            json.dump({'agents': agents_list}, f, indent=2)
        logger.info(f"Agents data saved to {AGENTS_JSON_PATH}")
    except Exception as e:
        logger.error(f"Error saving agents to {AGENTS_JSON_PATH}: {e}")

def load_agent_data():
    """Load agent data from JSON file."""
    global AGENTS_DATA
    AGENTS_DATA = {}
    if os.path.exists(AGENTS_JSON_PATH):
        try:
            with open(AGENTS_JSON_PATH, 'r', encoding='utf-8') as f:
                data = json.load(f)
                for agent in data.get('agents', []):
                    if 'pdfSources' not in agent and 'pdfSource' in agent:
                        agent['pdfSources'] = [agent['pdfSource']]
                    AGENTS_DATA[agent['agentId']] = agent
            logger.info(f"Loaded {len(AGENTS_DATA)} agents from {AGENTS_JSON_PATH}")
            return
        except Exception as e:
            logger.error(f"Error loading agent data from {AGENTS_JSON_PATH}: {e}")
    backup_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'agents.json')
    if os.path.exists(backup_path):
        try:
            with open(backup_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                for agent in data.get('agents', []):
                    if 'pdfSources' not in agent and 'pdfSource' in agent:
                        agent['pdfSources'] = [agent['pdfSource']]
                    AGENTS_DATA[agent['agentId']] = agent
            logger.info(f"Loaded {len(AGENTS_DATA)} agents from backup path {backup_path}")
            save_agents_to_json()
            return
        except Exception as e:
            logger.error(f"Error loading agent data from backup path {backup_path}: {e}")
    default_agents = [
        {
            "iconType": "FaShieldAlt",
            "name": "DPDP Compliance",
            "description": "Get insights from the Digital Personal Data Protection Act. Ask about user rights, data fiduciaries, and obligations.",
            "buttonText": "Start Chat",
            "agentId": "DPDP",
            "pdfSources": ["DPDP_act.pdf"]
        },
        {
            "iconType": "FaGavel",
            "name": "Parliamentary Rules",
            "description": "Access information on the Rules of Procedure and Conduct of Business in Lok Sabha.",
            "buttonText": "Start Chat",
            "agentId": "Parliament",
            "pdfSources": ["Rules_of_Procedures_Lok_Sabha.pdf"]
        }
    ]
    for agent in default_agents:
        AGENTS_DATA[agent['agentId']] = agent
    save_agents_to_json()
    try:
        with open(backup_path, 'w', encoding='utf-8') as f:
            json.dump({'agents': default_agents}, f, indent=2)
    except Exception as e:
        logger.error(f"Error saving default agents to backup path {backup_path}: {e}")
    logger.info(f"Created {len(AGENTS_DATA)} default agents")

# Call this function on startup
load_agent_data()

# --- Agent-specific retrievers ---
# Build a vectorstore and retriever for each agent's PDF(s)
agent_retrievers = {}
all_documents = []

# Load all PDFs and build a general retriever as fallback
pdf_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'backend', 'pdfs')
pdf_files = [f for f in os.listdir(pdf_dir) if f.endswith('.pdf')]
pdf_path_map = {os.path.basename(f): os.path.join(pdf_dir, f) for f in pdf_files}

# Helper to load a PDF as LangChain Document(s)
def load_pdf_as_documents(pdf_filename):
    pdf_dir = app.config['UPLOAD_FOLDER']
    pdf_path = os.path.join(pdf_dir, pdf_filename)
    if not os.path.exists(pdf_path):
        logger.error(f"PDF file {pdf_filename} not found in {pdf_dir}")
        return []
    try:
        reader = PdfReader(pdf_path)
        docs = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                docs.append(Document(page_content=text, metadata={"source": pdf_filename, "page": i+1}))
        return docs
    except Exception as e:
        logger.error(f"Error loading PDF {pdf_filename}: {e}")
        return []
def batch_embed_and_index(docs, embeddings, batch_size=500):
    from backend.utils import initialize_faiss_index
    all_docs = []
    for i in range(0, len(docs), batch_size):
        batch = docs[i:i+batch_size]
        # Embed the batch (but do not use precomputed_vectors)
        # vectors = embeddings.embed_documents([doc.page_content for doc in batch])
        all_docs.extend(batch)
    # Now build the FAISS index with all docs (let initialize_faiss_index compute embeddings)
    vectorstore = initialize_faiss_index(all_docs, embeddings)
    return vectorstore

# Build agent-specific retrievers (PATCH: build for all file types, not just PDFs)
def extract_structured(file_path):
    ext = file_path.lower().split('.')[-1]
    try:
        if ext == 'csv':
            df = pd.read_csv(file_path)
        elif ext == 'xlsx':
            df = pd.read_excel(file_path)
        elif ext == 'json':
            with open(file_path, 'r', encoding='utf-8') as f:
                data = pyjson.load(f)
            if isinstance(data, list):
                df = pd.json_normalize(data)
            else:
                df = pd.json_normalize([data])
        else:
            raise ValueError('Unsupported structured file')
        if df.empty:
            logger.warning(f"Structured file {file_path} is empty.")
            return []
        return df.to_dict(orient='records')
    except Exception as e:
        logger.error(f"Failed to extract structured data from {file_path}: {e}")
        return []

def extract_text(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def extract_pdf(file_path):
    extracted_text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ''
            if text.strip():
                extracted_text.append(text)
            else:
                # Fallback: Convert page to image and run Tesseract OCR
                pil_img = page.to_image(resolution=300).original
                img_gray = pil_img.convert('L')
                img_np = np.array(img_gray)
                img_bin = cv2.adaptiveThreshold(img_np, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 2)
                img_pil = Image.fromarray(img_bin)
                ocr_text = pytesseract.image_to_string(img_pil)
                extracted_text.append(ocr_text)
    return '\n'.join(extracted_text)

def preprocess_image_for_ocr(image):
    img_gray = image.convert('L')
    img_np = np.array(img_gray)
    img_bin = cv2.adaptiveThreshold(img_np, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 2)
    img_pil = Image.fromarray(img_bin)
    if img_pil.width < 600:
        scale = 600 / img_pil.width
        new_size = (int(img_pil.width * scale), int(img_pil.height * scale))
        img_pil = img_pil.resize(new_size, Image.ANTIALIAS)
    return img_pil

def extract_image(file_path, use_trocr=False):
    try:
        try:
            image = Image.open(file_path).convert('RGB')
            logger.info(f"Loaded image {file_path} with Pillow.")
        except Exception as pil_e:
            logger.warning(f"Pillow failed to open {file_path}: {pil_e}. Trying OpenCV...")
            img_cv = cv2.imread(file_path)
            if img_cv is None:
                logger.error(f"OpenCV failed to open {file_path} as image.")
                return ''
            image = Image.fromarray(cv2.cvtColor(img_cv, cv2.COLOR_BGR2RGB))
            logger.info(f"Loaded image {file_path} with OpenCV.")
        preprocessed = preprocess_image_for_ocr(image)
        if not use_trocr:
            try:
                text = pytesseract.image_to_string(preprocessed)
                logger.info(f"Extracted text from image {file_path} using Tesseract, length: {len(text)} chars.")
                return text
            except Exception as tess_e:
                logger.error(f"Tesseract OCR failed for {file_path}: {tess_e}")
                return ''
        else:
            # Add TrOCR logic here if needed
            return ''
    except Exception as e:
        logger.error(f"Failed to extract text from image {file_path}: {e}")
        return ''

def extract_content(file_path, file_type, audio_transcript=None):
    ext = file_path.lower().split('.')[-1]
    if file_type == 'audio':
        return audio_transcript or ''
    elif ext in ['csv', 'xlsx', 'json']:
        return extract_structured(file_path)
    elif ext in ['txt']:
        return extract_text(file_path)
    elif ext in ['pdf']:
        return extract_pdf(file_path)
    elif ext in ['jpg', 'jpeg', 'png', 'bmp', 'webp', 'tiff']:
        return extract_image(file_path, use_trocr=False)
    else:
        return None

def extract_audio_transcript(file_path):
    if not os.path.exists(file_path):
        logger.error(f"Audio file not found for transcription: {file_path}")
        return None
    if not WHISPER_AVAILABLE:
        logger.error("Whisper is not available for audio transcription.")
        return None
    try:
        result = whisper_model.transcribe(file_path)
        return result['text']
    except Exception as e:
        logger.error(f"Whisper transcription failed for {file_path}: {e}")
        return None


import docx2txt

for agent_id, agent in AGENTS_DATA.items():
    pdfs = agent.get('pdfSources', [])
    agent_docs = []
    for filename in pdfs:
        ext = filename.lower().split('.')[-1]
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        try:
            logger.info(f"[STARTUP] Processing file '{filename}' (type: {ext}) for retriever...")
            if ext == 'pdf':
                agent_docs.extend(load_pdf_as_documents(filename))
            elif ext in ['csv', 'xlsx', 'json']:
                content = extract_structured(file_path)
                if content:
                    logger.info(f"[STARTUP] Extracted structured data from {filename}, {len(content)} rows.")
                    import json as pyjson
                    from langchain.text_splitter import RecursiveCharacterTextSplitter
                    docs = []
                    for i, row in enumerate(content):
                        row_text = pyjson.dumps(row, ensure_ascii=False)
                        if len(row_text) > 1000:
                            chunks = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100).split_text(row_text)
                            for chunk in chunks:
                                docs.append(Document(page_content=chunk, metadata={"source": filename, "file_type": ext, "row": i}))
                        else:
                            docs.append(Document(page_content=row_text, metadata={"source": filename, "file_type": ext, "row": i}))
                    logger.info(f"[STARTUP] Created {len(docs)} document chunks for {filename}")
                    # Batch embed and index
                    agent_vectorstore = batch_embed_and_index(docs, embeddings, batch_size=500)
                    agent_retrievers[agent_id] = agent_vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 20})
                    logger.info(f"[PATCH] Built retriever for new agent {agent_id} with {len(docs)} docs (batched embedding).")
                    agent_docs.extend(docs)
                else:
                    logger.warning(f"[STARTUP] No structured data extracted from {filename} (type: {ext})")
                    continue
            elif ext in ['jpg', 'jpeg', 'png', 'bmp', 'webp', 'tiff']:
                content = extract_image(file_path, use_trocr=False)
                if content:
                    logger.info(f"[STARTUP] Extracted text from image {filename}, length: {len(content)} chars.")
                    from langchain.text_splitter import RecursiveCharacterTextSplitter
                    docs = [Document(page_content=chunk, metadata={"source": filename, "file_type": ext})
                            for chunk in RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100).split_text(content) if chunk.strip()]
                    logger.info(f"[STARTUP] Created {len(docs)} document chunks for {filename}")
                    agent_docs.extend(docs)
                else:
                    logger.warning(f"[STARTUP] No text extracted from image {filename} (type: {ext})")
            elif ext in ['mp3', 'wav', 'ogg', 'm4a']:
                transcript_path = os.path.splitext(file_path)[0] + '.txt'
                if os.path.exists(transcript_path):
                    with open(transcript_path, 'r', encoding='utf-8') as tf:
                        content = tf.read()
                    logger.info(f"[STARTUP] Loaded transcript for {filename}, length: {len(content)} chars")
                else:
                    # Try to generate transcript with Whisper
                    content = extract_audio_transcript(file_path)
                    if content:
                        with open(transcript_path, 'w', encoding='utf-8') as tf:
                            tf.write(content)
                        logger.info(f"Transcribed and saved transcript for {filename} at {transcript_path}")
                    else:
                        logger.warning(f"No transcript found or generated for audio file {filename}. Skipping retriever creation for this file.")
                        continue
                if content:
                    from langchain.text_splitter import RecursiveCharacterTextSplitter
                    docs = [Document(page_content=chunk, metadata={"source": filename, "file_type": ext})
                            for chunk in RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100).split_text(content) if chunk.strip()]
                    logger.info(f"Created {len(docs)} document chunks for {filename}")
                    agent_docs.extend(docs)
            elif ext == 'txt':
                content = extract_text(file_path)
                if content:
                    logger.info(f"[STARTUP] Extracted text from {filename}, length: {len(content)} chars.")
                    from langchain.text_splitter import RecursiveCharacterTextSplitter
                    docs = [Document(page_content=chunk, metadata={"source": filename, "file_type": ext})
                            for chunk in RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100).split_text(content) if chunk.strip()]
                    logger.info(f"[STARTUP] Created {len(docs)} document chunks for {filename}")
                    agent_docs.extend(docs)
                else:
                    logger.warning(f"[STARTUP] No text extracted from {filename} (type: {ext})")
            elif ext == 'docx':
                try:
                    text = docx2txt.process(file_path)
                    from langchain.text_splitter import RecursiveCharacterTextSplitter
                    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
                    chunks = splitter.split_text(text)
                    docs = [Document(page_content=chunk, metadata={"source": filename, "file_type": ext}) for chunk in chunks]
                    logger.info(f"[STARTUP] Created {len(docs)} document chunks for DOCX {filename}")
                    agent_vectorstore = batch_embed_and_index(docs, embeddings, batch_size=500)
                    agent_retrievers[agent_id] = agent_vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 20})
                    logger.info(f"[PATCH] Built retriever for new agent {agent_id} with {len(docs)} docs (batched embedding, DOCX).")
                    agent_docs.extend(docs)
                except Exception as docx_e:
                    logger.warning(f"[STARTUP] Failed to process DOCX {filename}: {docx_e}")
                    continue
            else:
                logger.warning(f"[STARTUP] Unsupported file type for retriever: {filename} (type: {ext})")
        except Exception as e:
            logger.warning(f"[STARTUP] Failed to process file {filename}: {e}")
            continue
    all_documents.extend(agent_docs)

# Build a general retriever for all documents as fallback
general_vectorstore = initialize_faiss_index(all_documents, embeddings) if all_documents else None
general_retriever = general_vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 20}) if general_vectorstore else None

ollama_client = OllamaClient()

@app.route('/sessions/<session_id>/messages', methods=['POST'])
def add_message(session_id):
    """Add a message to a session"""
    if session_id not in sessions:
        return jsonify({'error': 'Session not found'}), 404
    
    data = request.json
    # Always use the translated English message for all backend logic
    user_message = str(data.get('message', '')) if data.get('message') is not None else ''
    user_name = str(data.get('userName', 'User')) if data.get('userName') is not None else 'User'
    agent_id = data.get('agentId', None)
    source_from_frontend = data.get('source', data.get('pdfSource', None))
    original_message = data.get('original_message', user_message)  # Optionally store original message for display
    lang = data.get('lang', 'en-IN')  # Default to BCP-47 code
    model_name = data.get('model', 'gemini')
    
    logger.info(f"Processing user message: {user_message} for agent: {agent_id}, source: {source_from_frontend}")
    
    # Use agent-specific retriever if available, else fallback to general retriever
    retriever_to_use = agent_retrievers.get(agent_id) if agent_id in agent_retrievers else general_retriever
    if retriever_to_use is None:
        logger.error("No retriever available for this agent or general context.")
        return jsonify({'error': 'No retriever available.'}), 500
    all_docs = retriever_to_use.get_relevant_documents(user_message)
    logger.info(f"Retrieved {len(all_docs)} documents for agent {agent_id if agent_id else 'general'}.")
    # Filter by selected source if provided
    if source_from_frontend:
        # If the selected source is an audio file, map to its transcript txt file
        audio_exts = ['mp3', 'wav', 'ogg', 'm4a']
        ext = source_from_frontend.split('.')[-1].lower()
        if ext in audio_exts:
            transcript_name = source_from_frontend.rsplit('.', 1)[0] + '.txt'
            docs = [doc for doc in all_docs if doc.metadata.get('source') == transcript_name]
            logger.info(f"Filtered docs to transcript for audio source '{source_from_frontend}' (using '{transcript_name}'): {len(docs)} remaining.")
        else:
            docs = [doc for doc in all_docs if doc.metadata.get('source') == source_from_frontend]
            logger.info(f"Filtered docs to selected source '{source_from_frontend}': {len(docs)} remaining.")
    else:
        docs = all_docs
    logger.info(f"Documents returned: {[doc.metadata.get('source') for doc in docs]}")
    
    # Build context with source metadata for each chunk
    def format_source(doc):
        filename = doc.metadata.get('source', 'Unknown')
        page = doc.metadata.get('page', None)
        section = doc.metadata.get('section', None)
        meta = f"[Source: {filename}"
        if page is not None:
            meta += f", Page: {page}"
        if section is not None:
            meta += f", Section: {section}"
        meta += "]"
        return f"{meta}\n{doc.page_content}"
    context = "\n\n".join([format_source(doc) for doc in docs])
    # Escape curly braces in the context to prevent them from being misinterpreted as variables by Langchain
    escaped_context = context.replace('{', '{{').replace('}', '}}')
    logger.info(f"Retrieved context (first 500 chars): {escaped_context[:500]}...")
    
    # Dynamically determine document_heading based on retrieved documents
    current_doc_heading = "No Document Loaded"
    if docs:
        # Get unique sources from the retrieved documents
        unique_sources = list(set([doc.metadata.get('source', 'Unknown Document') for doc in docs]))
        if len(unique_sources) == 1:
            current_doc_heading = unique_sources[0]
        elif len(unique_sources) > 1:
            current_doc_heading = ", ".join(unique_sources) # Combine multiple sources

    # Get document ID from the first retrieved document's metadata (still useful for prompt metadata)
    document_id = docs[0].metadata.get('source', '') if docs else ''
    
    # Get the structured prompt, with extra explicit instruction for source links
    source_link_instruction = (
        "IMPORTANT: Always cite your sources using the exact filename and any available metadata as shown in the context below. "
        "For PDFs, use the format: pdf://<filename>/page/<page_number>#section=<section_name>. For other files, cite the filename and any available metadata. "
        "Do not invent filenames or page numbers. Do not add disclaimers, do not repeat instructions, and do not add extra commentary.\n\n"
    )
    prompt_data = create_structured_prompt(user_message, escaped_context, doc_id=document_id, DOCUMENT_HEADING=current_doc_heading, USER_NAME=user_name, AGENT_ID=agent_id)
    if 'prompt' in prompt_data:
        prompt_data['prompt'] = source_link_instruction + prompt_data['prompt']
    
    if model_name == 'llama3':
        from prompts import CORE_INSTRUCTION, FORMAT_PROMPT, DETAILED_RESPONSE_INSTRUCTION, FILTER_PROMPT, GREETING_INSTRUCTION, SUMMARIZATION_INSTRUCTION, COMPARISON_INSTRUCTION, EXTRACTION_INSTRUCTION, ANALYSIS_INSTRUCTION, TECHNICAL_INSTRUCTION, DEFINITION_INSTRUCTION, EXAMPLE_INSTRUCTION, TRANSLATION_INSTRUCTION, CLARIFICATION_INSTRUCTION, GENERAL_INSTRUCTION
        llama_system_prompt = f"""
You are a helpful AI assistant. Follow these instructions carefully:

{CORE_INSTRUCTION}

{FORMAT_PROMPT}

{DETAILED_RESPONSE_INSTRUCTION}

{FILTER_PROMPT}

{GREETING_INSTRUCTION}

{SUMMARIZATION_INSTRUCTION}

{COMPARISON_INSTRUCTION}

{EXTRACTION_INSTRUCTION}

{ANALYSIS_INSTRUCTION}

{TECHNICAL_INSTRUCTION}

{DEFINITION_INSTRUCTION}

{EXAMPLE_INSTRUCTION}

{TRANSLATION_INSTRUCTION}

{CLARIFICATION_INSTRUCTION}

{GENERAL_INSTRUCTION}

IMPORTANT: Only answer using the provided context. Do NOT use outside knowledge. Always cite your sources using the exact filename, page, and section as shown in the context. Use the format: pdf://<filename>/page/<page_number>#section=<section_name>. Do not invent filenames or page numbers. Do not add disclaimers, do not repeat instructions, and do not add extra commentary.

If any part of your answer is especially important, always underline it using double underscores (e.g., __this is important__). You may also use bold (**like this**) for key concepts, and both (**__like this__**) for the most critical points.
"""
        llama_user_message = f"""{context}

User's Question: {user_message}

REMEMBER: Start your answer with \"Thank you for asking. As per the information available to me:\", use **bold** and __underline__ as shown in the example, and end with a markdown Sources section as shown in the example above, using the actual filename, page, section, and a highlight from the context. Do NOT use [Source: ...] or any other format for sources.
"""
        response = ollama_client.chat(
            model='llama3',
            messages=[
                {"role": "system", "content": llama_system_prompt},
                {"role": "user", "content": llama_user_message}
            ]
        )['message']['content']
        model_used = 'llama3'
        # ENFORCEMENT: Only skip post-processing if a valid markdown link is present
        valid_filenames = set(doc.metadata.get('source') for doc in docs)
        found_valid_markdown_link = False
        markdown_link_pattern = re.compile(r'\(pdf://[\w\-.]+\.pdf(?:/page/\d+)?(?:#section=[^&\s)]+)?(?:&highlight=[^\s)]+)?\)')
        if markdown_link_pattern.search(response):
            found_valid_markdown_link = True
        if not found_valid_markdown_link:
            logger.warning(f"[LLAMA ENFORCEMENT] No valid markdown link found in response. Valid filenames: {valid_filenames}")
            answer_text = response.strip()
            # Only use the top relevant chunk for the fallback link
            if docs:
                top_doc = docs[0]
                filename = top_doc.metadata.get('source', 'Unknown.pdf')
                page = top_doc.metadata.get('page', None)
                section = top_doc.metadata.get('section', None)
                # Use up to 200 characters of the chunk as the highlight
                words = top_doc.page_content[:200]
                highlight = words.replace('\n', ' ')
                link = f"(pdf://{filename}"
                if page:
                    link += f"/page/{page}"
                if section:
                    link += f"#section={section}"
                if highlight:
                    from urllib.parse import quote
                    link += f"&highlight={quote(highlight)}"
                link += ")"
                # Remove any line that starts with (optionally bolded) 'Sources', with or without a colon, and any text after it
                answer_text = re.sub(r'(?im)^\s*(\*\*\s*)?Sources\*\*?\s*:?.*$', '', answer_text)
                response = f"{answer_text.strip()}\n\n**Sources:**\n- {link}"
            else:
                response = "I cannot answer based on the provided information."
        # Log the final response string before returning
        logger.info(f"[FINAL LLAMA RESPONSE TO FRONTEND] {response}")
        # Remove all but the last 'Sources:' section (optionally bolded, with or without colon)
        # Split by lines, keep only the last occurrence of a 'Sources:' line and its following lines
        lines = response.splitlines()
        sources_indices = [i for i, line in enumerate(lines) if re.match(r'(?i)^\s*(\*\*\s*)?Sources\*\*?\s*:?.*$', line)]
        if len(sources_indices) > 1:
            # Remove all but the last 'Sources:' line and its following lines
            last_idx = sources_indices[-1]
            # Keep everything up to the last 'Sources:' line, then only the last 'Sources:' section
            lines = lines[:last_idx] + lines[last_idx:]
            response = '\n'.join(lines)
        # After getting the response, parse plain text 'Sources:' block and convert to markdown links
        sources_block_match = re.search(r'Sources?:\s*(.*)', response, re.IGNORECASE | re.DOTALL)
        markdown_links = []
        if sources_block_match:
            sources_text = sources_block_match.group(1)
            response = response[:sources_block_match.start()].strip()
            for match in re.finditer(r'\(([^)]+)\),?\s*Page:\s*(\d+)', sources_text):
                section = match.group(1).strip()
                page = match.group(2).strip()
                # Try to find the correct filename from docs
                filename = docs[0].metadata.get('source', 'Unknown.pdf') if docs else 'Unknown.pdf'
                link = f"- (pdf://{filename}/page/{page}#section={section.replace(' ', '_')})"
                markdown_links.append(link)
        # If we found markdown links, append them as the sources section
        if markdown_links:
            response += "\n\n**Sources:**\n" + "\n".join(markdown_links)
    elif model_name == 'mistral':
        from prompts import CORE_INSTRUCTION, FORMAT_PROMPT, DETAILED_RESPONSE_INSTRUCTION, FILTER_PROMPT, GREETING_INSTRUCTION, SUMMARIZATION_INSTRUCTION, COMPARISON_INSTRUCTION, EXTRACTION_INSTRUCTION, ANALYSIS_INSTRUCTION, TECHNICAL_INSTRUCTION, DEFINITION_INSTRUCTION, EXAMPLE_INSTRUCTION, TRANSLATION_INSTRUCTION, CLARIFICATION_INSTRUCTION, GENERAL_INSTRUCTION
        mistral_system_prompt = f"""
You are a helpful AI assistant. Follow these instructions carefully:

- Use **bold** for key concepts.
- Use __underline__ for especially important parts.
- Use both (**__like this__**) for the most critical points.
- Only answer using the provided context. Do NOT use outside knowledge.
- Always cite your sources using the exact filename, page, and section as shown in the context. Use the format: pdf://<filename>/page/<page_number>#section=<section_name>. Do not invent filenames or page numbers. Do not add disclaimers, do not repeat instructions, and do not add extra commentary.
- If the answer is not found in the context, state that you cannot answer based on the provided information.
"""
        mistral_user_message = f"""{context}

User's Question: {user_message}

REMEMBER: Start your answer with \"Thank you for asking. As per the information available to me:\", use **bold** and __underline__ as shown in the example, and end with a markdown Sources section as shown in the example above, using the actual filename, page, section, and a highlight from the context. Do NOT use [Source: ...] or any other format for sources.
"""
        response = ollama_client.chat(
            model='mistral',
            messages=[
                {"role": "system", "content": mistral_system_prompt},
                {"role": "user", "content": mistral_user_message}
            ]
        )['message']['content']
        model_used = 'mistral'
        # ENFORCEMENT: Only skip post-processing if a valid markdown link is present
        valid_filenames = set(doc.metadata.get('source') for doc in docs)
        found_valid_markdown_link = False
        markdown_link_pattern = re.compile(r'\(pdf://[\w\-.]+\.pdf(?:/page/\d+)?(?:#section=[^&\s)]+)?(?:&highlight=[^\s)]+)?\)')
        if markdown_link_pattern.search(response):
            found_valid_markdown_link = True
        if not found_valid_markdown_link:
            logger.warning(f"[MISTRAL ENFORCEMENT] No valid markdown link found in response. Valid filenames: {valid_filenames}")
            answer_text = response.strip()
            # Ultra-flexible regex to match various citation patterns
            source_pattern = re.compile(r"""
                (?:
                    (?P<filename>[\w\-]+\.pdf)[\s,;:()\n]*
                    (?:Section[:\s]*?(?P<section1>[\w\-]+)[\s,;:()\n]*)?
                    (?:Page[:\s]*?(?P<page1>\d+))?
                )|
                (?:
                    (?:Page[:\s]*?(?P<page2>\d+)[\s,;:()\n]*)
                    (?:Section[:\s]*?(?P<section2>[\w\-]+)[\s,;:()\n]*)?
                    (?P<filename2>[\w\-]+\.pdf)
                )|
                (?:
                    (?:Section[:\s]*?(?P<section3>[\w\-]+)[\s,;:()\n]*)
                    (?P<filename3>[\w\-]+\.pdf)[\s,;:()\n]*
                    (?:Page[:\s]*?(?P<page3>\d+))?
                )|
                (?:
                    (?P<filename4>[\w\-]+\.pdf)[\s,;:()\n]*\(\s*Page\s*(?P<page4>\d+)\s*\)
                )|
                (?:
                    Page\s*(?P<page5>\d+)\s*of\s*(?P<filename5>[\w\-]+\.pdf)
                )
            """, re.IGNORECASE | re.VERBOSE)
            matches = list(source_pattern.finditer(answer_text))
            links = []
            spans_to_remove = []
            for match in matches:
                filename = (
                    match.group('filename') or match.group('filename2') or match.group('filename3') or match.group('filename4') or match.group('filename5')
                )
                page = (
                    match.group('page1') or match.group('page2') or match.group('page3') or match.group('page4') or match.group('page5')
                )
                section = (
                    match.group('section1') or match.group('section2') or match.group('section3')
                )
                if filename:
                    # Use a placeholder highlight since doc is not available in this context
                    highlight = "Document content preview"
                    link = f"(pdf://{filename}"
                    if page:
                        link += f"/page/{page}"
                    if section:
                        link += f"#section={section}"
                    if highlight:
                        from urllib.parse import quote
                        link += f"&highlight={quote(highlight)}"
                    link += ")"
                    links.append(f"- {link}")
                    # Record the span to remove
                    spans_to_remove.append(match.span())
            # Remove all matched citation spans from the answer text
            if spans_to_remove:
                # Remove from end to start to avoid messing up indices
                for start, end in sorted(spans_to_remove, reverse=True):
                    answer_text = answer_text[:start] + answer_text[end:]
            if links:
                # Remove any line that starts with (optionally bolded) 'Sources', with or without a colon, and any text after it
                answer_text = re.sub(r'(?im)^\s*(\*\*\s*)?Sources\*\*?\s*:?.*$', '', answer_text)
                response = f"{answer_text.strip()}\n\n**Sources:**\n" + '\n'.join(links)
            else:
                if docs:
                    top_doc = docs[0]
                    filename = top_doc.metadata.get('source', 'Unknown.pdf')
                    page = top_doc.metadata.get('page', None)
                    section = top_doc.metadata.get('section', None)
                    words = top_doc.page_content.split()
                    highlight = ' '.join(words[:20]) if words else ''
                    link = f"(pdf://{filename}"
                    if page:
                        link += f"/page/{page}"
                    if section:
                        link += f"#section={section}"
                    if highlight:
                        from urllib.parse import quote
                        link += f"&highlight={quote(highlight)}"
                    link += ")"
                    answer_text = re.sub(r'(?im)^\s*(\*\*\s*)?Sources\*\*?\s*:?.*$', '', answer_text)
                    response = f"{answer_text.strip()}\n\n**Sources:**\n- {link}"
                else:
                    response = "I cannot answer based on the provided information."
        # Log the final response string before returning
        logger.info(f"[FINAL MISTRAL RESPONSE TO FRONTEND] {response}")
        # Remove all but the last 'Sources:' section (optionally bolded, with or without colon)
        # Split by lines, keep only the last occurrence of a 'Sources:' line and its following lines
        lines = response.splitlines()
        sources_indices = [i for i, line in enumerate(lines) if re.match(r'(?i)^\s*(\*\*\s*)?Sources\*\*?\s*:?.*$', line)]
        if len(sources_indices) > 1:
            # Remove all but the last 'Sources:' line and its following lines
            last_idx = sources_indices[-1]
            # Keep everything up to the last 'Sources:' line, then only the last 'Sources:' section
            lines = lines[:last_idx] + lines[last_idx:]
            response = '\n'.join(lines)
    else:
        # Default: Gemini
        # Create the chain
        chain = (
            {"context": RunnablePassthrough(), "user_message": RunnablePassthrough()}
            | ChatPromptTemplate.from_template(prompt_data["prompt"])
            | llm
            | StrOutputParser()
        )
        # Generate response
        response = chain.invoke({
            "context": escaped_context,
            "user_message": user_message
        })
        model_used = 'gemini'
        # Debug: log docs for Gemini
        print(f"[DEBUG] Gemini docs: {[doc.page_content[:100] for doc in docs]}")
    
    logger.info(f"Raw model response (first 500 chars): {response[:500]}...")
    
    # Format the response with source information
    formatted_response = format_response_with_sources(response, docs)
    # Safety net: replace any accidental '@pdf://' with 'pdf://' in the final response
    if formatted_response:
        formatted_response = formatted_response.replace('@pdf://', 'pdf://')

    # Remove all but the last 'Sources:' section (optionally bolded, with or without colon)
    def remove_duplicate_sources_sections(text):
        pattern = re.compile(r'(?im)^\s*(\*{0,2}\s*)?Sources(\*{0,2})?\s*:?.*$', re.MULTILINE)
        matches = list(pattern.finditer(text))
        if len(matches) <= 1:
            return text
        last_idx = matches[-1].start()
        return text[:last_idx] + text[last_idx:]
    formatted_response = remove_duplicate_sources_sections(formatted_response)
    
    logger.info(f"Formatted sources after processing: {formatted_response.split('**Sources:**')[1] if '**Sources:**' in formatted_response else 'No sources block'}")
    
    # Log the formatted response (first 500 chars)
    if formatted_response is None:
        logger.error("Received None response from format_response_with_sources")
        formatted_response = "I apologize, but I encountered an error while processing your request. Please try again."
    else:
        logger.info(f"Formatted response (first 500 chars): {formatted_response[:500]}...")
    
    # Add messages to session
    sessions[session_id]['messages'].append({
        'sender': 'user',
        'content': original_message,  # Store the original message for display
        'agentId': agent_id,  # Store the agent ID received from frontend or inferred
        'lang': lang  # Store BCP-47 code
    })
    sessions[session_id]['messages'].append({
        'sender': 'assistant',
        'content': str(formatted_response), # Ensure assistant response is always a string
        'agentId': agent_id, # Store the agent ID for assistant messages as well
        'lang': lang,  # Store BCP-47 code
        'model_used': model_used  # Store the model used for this response
    })
    
    # Log monitoring data
    try:
        # Get user email from request or use a default
        user_email = data.get('userEmail', 'anonymous@user.com')
        
        # Log model usage
        log_model_usage(
            user_email=user_email,
            model_name=model_used,
            session_id=session_id,
            agent_id=agent_id,
            response_time=None,  # Could be calculated if we track timing
            tokens_used=None,    # Could be extracted from response if available
            success=True
        )
        
        # Log user activity
        log_user_activity(
            user_email=user_email,
            action='send_message',
            session_id=session_id,
            agent_id=agent_id,
            details={'message_length': len(user_message)}
        )
        
        # Log agent usage if agent is specified
        if agent_id:
            agent_name = "Unknown Agent"
            # Try to get agent name from agent data
            try:
                # Use the global AGENTS_DATA instead of calling load_agent_data()
                if agent_id in AGENTS_DATA:
                    agent_name = AGENTS_DATA[agent_id].get('name', 'Unknown Agent')
            except Exception as e:
                logger.error(f"Error getting agent name for {agent_id}: {e}")
            
            log_agent_usage(
                agent_id=agent_id,
                agent_name=agent_name,
                user_email=user_email,
                session_id=session_id,
                usage_type='interaction'
            )
    except Exception as e:
        logger.error(f"Error logging monitoring data: {e}")
    
    # Update session title if it's the first message
    if len(sessions[session_id]['messages']) == 2:
        # Generate title based on the first user message
        first_user_message_content = sessions[session_id]['messages'][0]['content']
        # Try to extract the first sentence
        first_sentence_match = re.match(r'([^.!?\n]+[.!?\n]?)', first_user_message_content)
        if first_sentence_match:
            generated_title = first_sentence_match.group(1).strip()
        else:
            # Fallback to the beginning of the message if no sentence ending punctuation
            generated_title = first_user_message_content.strip()

        # Truncate to a reasonable length (e.g., 50 characters) and add ellipsis if truncated
        max_title_length = 50
        if len(generated_title) > max_title_length:
            sessions[session_id]['title'] = generated_title[:max_title_length].strip() + '...'
        elif generated_title:
            sessions[session_id]['title'] = generated_title
        else:
            # Fallback if somehow the message was empty after stripping
            sessions[session_id]['title'] = f'Chat {len(sessions)}'

    save_sessions() # Save sessions after adding a message or updating title
    
    # After docs are retrieved, get source highlights
    source_highlights = [doc.page_content for doc in docs if doc.page_content]
    
    # Remove markdown source links from main answer text before sending to frontend
    if formatted_response and '**Sources:**' in formatted_response:
        main_answer, sources_section = formatted_response.split('**Sources:**', 1)
        main_answer = re.sub(r'\(pdf://[^\)]+\)', '', main_answer).strip()
        formatted_response = f"{main_answer}\n\n**Sources:**{sources_section}"
    
    return jsonify({
        'response': formatted_response,
        'query_type': prompt_data['query_type'],
        'metadata': prompt_data['metadata'],
        'session': sessions[session_id],
        'agentId': agent_id,
        'model_used': model_used,
        'source_highlights': source_highlights,
        'sources': [format_source(doc) for doc in docs]
    })

# Add route to serve PDF files
@app.route('/pdfs/<path:filename>')
def serve_pdf(filename):
    """Serve files from the backend/pdfs directory with correct MIME type"""
    try:
        pdf_directory = app.config['UPLOAD_FOLDER']
        pdf_path = os.path.join(pdf_directory, filename)
        if not os.path.exists(pdf_path):
            logger.error(f"File not found: {pdf_path}")
            return jsonify({'error': 'File not found'}), 404
        # Guess the correct MIME type
        mime_type, _ = mimetypes.guess_type(pdf_path)
        if not mime_type:
            mime_type = 'application/octet-stream'
        with open(pdf_path, 'rb') as f:
            file_data = f.read()
        response = app.response_class(
            response=file_data,
            status=200,
            mimetype=mime_type
        )
        # Add CORS headers
        response.headers['Access-Control-Allow-Origin'] = 'http://localhost:3000'
        response.headers['Access-Control-Allow-Methods'] = 'GET, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        response.headers['Access-Control-Allow-Credentials'] = 'true'
        return response
    except Exception as e:
        logger.error(f"Error serving file: {str(e)}")
        return jsonify({'error': str(e)}), 500

ALLOWED_EXTENSIONS = {'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Add this function to rebuild the FAISS index from all PDFs in the upload folder
def rebuild_faiss_index():
    global retriever
    pdf_documents = load_pdfs_from_directory(app.config['UPLOAD_FOLDER'])
    vectorstore = initialize_faiss_index(pdf_documents, embeddings)
    retriever = vectorstore.as_retriever(
        search_type="similarity",
        search_kwargs={"k": 20}
    )

@app.route('/upload-pdf', methods=['POST', 'OPTIONS'])
def upload_pdf():
    """Handle PDF file uploads (multiple files supported)"""
    if request.method == 'OPTIONS':
        return '', 200

    try:
        if 'files' not in request.files:
            logger.error("No files part in request")
            return jsonify({'error': 'No files part', 'success': False}), 400
        
        files = request.files.getlist('files')
        if not files or all(f.filename == '' for f in files):
            logger.error("No selected files")
            return jsonify({'error': 'No selected files', 'success': False}), 400
        
        saved_filenames = []
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
                logger.info(f"Saving file: {filepath}")
                file.save(filepath)
                saved_filenames.append(filename)
            else:
                logger.error(f"Invalid file type: {file.filename}")
                return jsonify({'error': f'Invalid file type: {file.filename}', 'success': False}), 400
        try:
            logger.info(f"Rebuilding FAISS index after uploading: {saved_filenames}")
            rebuild_faiss_index()
            logger.info(f"FAISS index rebuilt successfully after uploading: {saved_filenames}")
            return jsonify({
                'message': 'Files uploaded and processed successfully',
                'filenames': saved_filenames,
                'success': True
            })
        except Exception as e:
            logger.error(f"Error during FAISS index rebuild: {str(e)}")
            # Remove all uploaded files if processing fails
            for filename in saved_filenames:
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                if os.path.exists(filepath):
                    os.remove(filepath)
            return jsonify({'error': f'Error processing PDFs: {str(e)}', 'success': False}), 500
    except Exception as e:
        logger.error(f"Unexpected error in upload_pdf route: {str(e)}")
        return jsonify({'error': f'Unexpected error: {str(e)}', 'success': False}), 500

# Load PDFs from uploads directory on startup
def load_uploaded_pdfs():
    global retriever # Declare retriever as global to modify it
    existing_pdfs_in_index = set()

    # Try to load existing FAISS index to get sources
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")  # Multilingual support
    if os.path.exists("faiss_index"):
        try:
            vectorstore = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
            # Extract existing sources from metadata (this might be tricky without iterating all docs)
            # For simplicity, we'll assume if it's in the folder, we should try to process it
            retriever = vectorstore.as_retriever()
            logger.info("Existing FAISS index loaded.")
        except Exception as e:
            logger.warning(f"Error loading existing FAISS index: {e}. A new index will be created if needed.")
            retriever = None # Reset retriever if loading fails
    else:
        logger.info("No existing FAISS index found. A new one will be created.")
        retriever = None

    # Process all PDFs in the UPLOAD_FOLDER
    if os.path.exists(UPLOAD_FOLDER):
        for filename in os.listdir(UPLOAD_FOLDER):
            if filename.endswith('.pdf'):
                filepath = os.path.join(UPLOAD_FOLDER, filename)
                if retriever is None or filename not in existing_pdfs_in_index: # Only process if new or retriever is not initialized
                    try:
                        process_pdf(filepath) # This will update the global retriever
                        logger.info(f"Loaded PDF: {filename}")
                    except Exception as e:
                        logger.error(f"Error loading PDF {filename}: {str(e)}")
    
    if retriever is None:
        logger.warning("No PDFs processed and no retriever initialized. Chat functionality may be limited.")

# Call this function on startup
load_uploaded_pdfs()

# Initialize global retriever after load_uploaded_pdfs has potentially set it
# If no PDFs were found, `retriever` will remain None or be a mock.
if retriever is None:
    logger.warning("Retriever is still None after loading PDFs. Initializing a mock retriever.")
    class FallbackRetriever:
        def get_relevant_documents(self, query):
            logger.warning("Using FallbackRetriever: No documents indexed. Query will not have context.")
            return []
    retriever = FallbackRetriever()

# Agent management endpoints
@app.route('/agents', methods=['GET'])
def get_agents():
    """Get all agents from AGENTS_DATA"""
    try:
        # Reload agents from file to ensure we have the latest data
        load_agent_data()
        agents_list = list(AGENTS_DATA.values())
        logger.info(f"GET /agents: Returning {len(agents_list)} agents")
        return jsonify(agents_list)
    except Exception as e:
        logger.error(f"Error in get_agents: {e}")
        return jsonify([])

@app.route('/agents', methods=['POST'])
def add_agent():
    """Add a new agent to AGENTS_DATA and save to agents.json (supports multiple pdfSources)"""
    try:
        data = request.json
        agent_id = data.get('agentId')
        if not agent_id:
            return jsonify({'error': 'agentId is required', 'success': False}), 400
        
        if agent_id in AGENTS_DATA:
            return jsonify({'error': f'Agent with ID {agent_id} already exists', 'success': False}), 409

        # Accept sources (list) or pdfSources/pdfSource for backward compatibility
        sources = data.get('sources')
        if not sources:
            pdf_sources = data.get('pdfSources')
            if not pdf_sources:
                pdf_source = data.get('pdfSource')
                if pdf_source:
                    pdf_sources = [pdf_source]
                else:
                    return jsonify({'error': 'Missing required agent sources/pdfSources/pdfSource', 'success': False}), 400
            sources = pdf_sources

        # Ensure required fields are present
        required_fields = ['name', 'description', 'iconType']
        if not all(field in data for field in required_fields):
            return jsonify({'error': 'Missing required agent fields', 'success': False}), 400

        agent_data = {
            'agentId': agent_id,
            'name': data['name'],
            'description': data['description'],
            'iconType': data['iconType'],
            'buttonText': data.get('buttonText', 'Start Chat'),
            'sources': sources,
            'tileLineStartColor': data.get('tileLineStartColor', ''),
            'tileLineEndColor': data.get('tileLineEndColor', ''),
        }
        AGENTS_DATA[agent_id] = agent_data
        save_agents_to_json()
        logger.info(f"Added new agent: {agent_id}")

        # --- PATCH START: Build retriever for new agent ---
        # Load PDFs and build retriever for the new agent
        pdfs = agent_data.get('pdfSources', [])
        agent_docs = []
        for filename in pdfs:
            ext = filename.lower().split('.')[-1]
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            try:
                logger.info(f"Processing file '{filename}' (type: {ext}) for retriever...")
                if ext == 'pdf':
                    agent_docs.extend(load_pdf_as_documents(filename))
                elif ext in ['csv', 'xlsx', 'json']:
                    content = extract_structured(file_path)
                    if content:
                        logger.info(f"Extracted structured data from {filename}, {len(content)} rows.")
                        import json as pyjson
                        from langchain.text_splitter import RecursiveCharacterTextSplitter
                        docs = []
                        for i, row in enumerate(content):
                            row_text = pyjson.dumps(row, ensure_ascii=False)
                            if len(row_text) > 1000:
                                chunks = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100).split_text(row_text)
                                for chunk in chunks:
                                    docs.append(Document(page_content=chunk, metadata={"source": filename, "file_type": ext, "row": i}))
                            else:
                                docs.append(Document(page_content=row_text, metadata={"source": filename, "file_type": ext, "row": i}))
                        logger.info(f"Created {len(docs)} document chunks for {filename}")
                        # Batch embed and index
                        agent_vectorstore = batch_embed_and_index(docs, embeddings, batch_size=500)
                        agent_retrievers[agent_id] = agent_vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 20})
                        logger.info(f"[PATCH] Built retriever for new agent {agent_id} with {len(docs)} docs (batched embedding).")
                        agent_docs.extend(docs)
                    else:
                        logger.warning(f"No structured data extracted from {filename} (type: {ext})")
                        continue
                elif ext in ['jpg', 'jpeg', 'png', 'bmp', 'webp', 'tiff']:
                    content = extract_image(file_path, use_trocr=False)
                    if content:
                        logger.info(f"Extracted text from image {filename}, length: {len(content)} chars.")
                        from langchain.text_splitter import RecursiveCharacterTextSplitter
                        docs = [Document(page_content=chunk, metadata={"source": filename, "file_type": ext})
                                for chunk in RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100).split_text(content) if chunk.strip()]
                        logger.info(f"Created {len(docs)} document chunks for {filename}")
                        agent_docs.extend(docs)
                    else:
                        logger.warning(f"No text extracted from image {filename} (type: {ext})")
                elif ext in ['mp3', 'wav', 'ogg', 'm4a']:
                    transcript_path = os.path.splitext(file_path)[0] + '.txt'
                    if os.path.exists(transcript_path):
                        with open(transcript_path, 'r', encoding='utf-8') as tf:
                            content = tf.read()
                        logger.info(f"Loaded transcript for {filename}, length: {len(content)} chars")
                    else:
                        # Try to generate transcript with Whisper
                        content = extract_audio_transcript(file_path)
                        if content:
                            with open(transcript_path, 'w', encoding='utf-8') as tf:
                                tf.write(content)
                            logger.info(f"Transcribed and saved transcript for {filename} at {transcript_path}")
                        else:
                            logger.warning(f"No transcript found or generated for audio file {filename}. Skipping retriever creation for this file.")
                            continue
                    if content:
                        from langchain.text_splitter import RecursiveCharacterTextSplitter
                        docs = [Document(page_content=chunk, metadata={"source": filename, "file_type": ext})
                                for chunk in RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100).split_text(content) if chunk.strip()]
                        logger.info(f"Created {len(docs)} document chunks for {filename}")
                        agent_docs.extend(docs)
                elif ext == 'txt':
                    content = extract_text(file_path)
                    if content:
                        logger.info(f"Extracted text from {filename}, length: {len(content)} chars.")
                        from langchain.text_splitter import RecursiveCharacterTextSplitter
                        docs = [Document(page_content=chunk, metadata={"source": filename, "file_type": ext})
                                for chunk in RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100).split_text(content) if chunk.strip()]
                        logger.info(f"Created {len(docs)} document chunks for {filename}")
                        agent_docs.extend(docs)
                    else:
                        logger.warning(f"No text extracted from {filename} (type: {ext})")
                elif ext == 'docx':
                    try:
                        text = docx2txt.process(file_path)
                        from langchain.text_splitter import RecursiveCharacterTextSplitter
                        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
                        chunks = splitter.split_text(text)
                        docs = [Document(page_content=chunk, metadata={"source": filename, "file_type": ext}) for chunk in chunks]
                        logger.info(f"[STARTUP] Created {len(docs)} document chunks for DOCX {filename}")
                        agent_vectorstore = batch_embed_and_index(docs, embeddings, batch_size=500)
                        agent_retrievers[agent_id] = agent_vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 20})
                        logger.info(f"[PATCH] Built retriever for new agent {agent_id} with {len(docs)} docs (batched embedding, DOCX).")
                        agent_docs.extend(docs)
                    except Exception as docx_e:
                        logger.warning(f"[STARTUP] Failed to process DOCX {filename}: {docx_e}")
                        continue
                else:
                    logger.warning(f"Unsupported file type for retriever: {filename} (type: {ext})")
            except Exception as e:
                logger.error(f"Error extracting/processing {filename} (type: {ext}): {e}")
            if agent_docs:
                agent_vectorstore = initialize_faiss_index(agent_docs, embeddings)
                agent_retrievers[agent_id] = agent_vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 20})
                logger.info(f"Built retriever for agent {agent_id} with {len(agent_docs)} docs.")
            else:
                logger.warning(f"[PATCH] No documents found for new agent {agent_id}. Retriever not built.")
            return jsonify({
                'message': 'Agent created successfully',
                'agent': agent_data,
                'success': True
            }), 201
    except Exception as e:
        logger.error(f"Error adding agent: {str(e)}")
        return jsonify({'error': f'Error adding agent: {str(e)}', 'success': False}), 500

@app.route('/agents/<agent_id>', methods=['PUT', 'OPTIONS'])
def update_agent(agent_id):
    """Update an existing agent"""
    if request.method == 'OPTIONS':
        return '', 200

    if agent_id not in AGENTS_DATA:
        return jsonify({'error': 'Agent not found'}), 404
    
    # Update only allowed fields, or merge entire object
    data = request.json
    AGENTS_DATA[agent_id].update(data) # Update the agent data
    save_agents_to_json() # Save changes to JSON file
    logger.info(f"Updated agent: {agent_id}")
    return jsonify(AGENTS_DATA[agent_id])

@app.route('/agents/<agent_id>', methods=['DELETE', 'OPTIONS'])
def delete_agent(agent_id):
    """Delete an agent from AGENTS_DATA and save to agents.json"""
    if request.method == 'OPTIONS':
        return '', 200

    try:
        if agent_id not in AGENTS_DATA:
            return jsonify({'error': 'Agent not found'}), 404
        
        # Get all PDF filenames before deleting the agent
        pdf_filenames = AGENTS_DATA[agent_id].get('pdfSources', [])
        
        # Delete the agent from AGENTS_DATA
        del AGENTS_DATA[agent_id]
        
        # Save to JSON file immediately after deletion
        save_agents_to_json()
        
        # Delete all associated PDF files if they exist
        for pdf_filename in pdf_filenames:
            pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
            try:
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                    logger.info(f"Deleted PDF file: {pdf_filename}")
            except Exception as e:
                logger.error(f"Error deleting PDF file {pdf_filename}: {str(e)}")
        
        logger.info(f"Deleted agent: {agent_id}")
        return '', 204
    except Exception as e:
        logger.error(f"Error deleting agent {agent_id}: {str(e)}")
        return jsonify({'error': f'Error deleting agent: {str(e)}'}), 500

@app.route('/elevenlabs/tts', methods=['POST', 'OPTIONS'])
def elevenlabs_tts():
    if request.method == 'OPTIONS':
        return '', 200
    try:
        data = request.json
        text = data.get('text')
        voice_lang = data.get('voiceLang', 'en-US') # Get the voice language from the request
        
        # Set voice_id based on language
        if voice_lang == 'hi-IN':
            voice_id ='TXb68m09B5U6BTh8UMd5'# Hindi voice ID
        else:
            voice_id = 'NFG5qt843uXKj4pFvR7C'  # Default English voice ID (Adam) for other languages

        if not text:
            return jsonify({'error': 'No text provided'}), 400

        # Use elevenlabs_client to generate audio
        audio_stream = elevenlabs_client.text_to_speech.stream(
            text=text,
            voice_id=voice_id,
            model_id="eleven_multilingual_v2", # Using a multilingual model
        )

        def generate_audio():
            for chunk in audio_stream:
                yield chunk

        response = app.response_class(generate_audio(), mimetype='audio/mpeg')
        response.headers['Content-Disposition'] = 'inline; filename="speech.mp3"'
        response.headers['Access-Control-Allow-Origin'] = 'http://localhost:3000'
        response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        response.headers['Access-Control-Allow-Credentials'] = 'true'
        return response

    except Exception as e:
        logger.error(f"ElevenLabs TTS error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/elevenlabs/stt', methods=['POST', 'OPTIONS'])
def elevenlabs_stt():
    if request.method == 'OPTIONS':
        return '', 200
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400

        audio_file = request.files['audio']
        import tempfile, os
        filename = audio_file.filename
        ext = os.path.splitext(filename)[1] or '.webm'
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            audio_file.save(tmp.name)
            tmp.flush()
            temp_path = tmp.name

        if not WHISPER_AVAILABLE:
            os.remove(temp_path)
            return jsonify({'error': 'Whisper is not available for audio transcription.'}), 500
        try:
            result = whisper_model.transcribe(temp_path, language=None)  # auto-detect language
            transcript = result['text']
            detected_language = result.get('language', 'unknown')
        except Exception as e:
            logger.error(f"Whisper transcription failed for uploaded audio: {e}")
            os.remove(temp_path)
            return jsonify({'error': str(e)}), 500

        os.remove(temp_path)
        response = jsonify({'transcript': transcript, 'language': detected_language})
        response.headers['Access-Control-Allow-Origin'] = 'http://localhost:3000'
        response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        response.headers['Access-Control-Allow-Credentials'] = 'true'
        return response

    except Exception as e:
        logger.error(f"ElevenLabs STT error: {e}")
        return jsonify({'error': str(e)}), 500

# --- Translation Consistency Helper (for future use) ---
def sarvam_translate_consistent(text, target_language_code, source_language_code="auto"):
    """
    Use Sarvam Translate with a prompt for legal/official queries to improve translation consistency.
    Accepts BCP-47 codes for language.
    """
    import requests
    SARVAM_API_KEY = os.getenv("SARVAM_API_KEY", "sk_y9798qpa_JJoiaSJY4GWdgVWW8Gs0LGN5")
    SARVAM_URL = "https://api.sarvam.ai/translate"
    headers = {
        "api-subscription-key": SARVAM_API_KEY,
        "Content-Type": "application/json"
    }
    payload = {
        "input": text,
        "source_language_code": source_language_code,
        "target_language_code": target_language_code
    }
    response = requests.post(SARVAM_URL, headers=headers, json=payload)
    if response.status_code == 200:
        return response.json().get("translated_text")
    else:
        return text  # fallback

@app.route('/sessions/clear_all', methods=['POST'])
def clear_all_sessions():
    """Delete all sessions and save to sessions.json"""
    global sessions
    sessions = {}
    save_sessions()
    return '', 204

PROJECTS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'backend', 'projects.json')
projects_lock = threading.Lock()

def load_projects():
    if not os.path.exists(PROJECTS_FILE):
        return []
    with open(PROJECTS_FILE, 'r', encoding='utf-8') as f:
        try:
            return json.load(f)
        except Exception:
            return []

def save_projects(projects):
    with projects_lock:
        with open(PROJECTS_FILE, 'w', encoding='utf-8') as f:
            json.dump(projects, f, indent=2)

@app.route('/projects', methods=['GET'])
def get_projects():
    projects = load_projects()
    return jsonify(projects)

@app.route('/projects', methods=['POST'])
def create_project():
    data = request.json
    projects = load_projects()
    new_project = {
        'id': str(uuid.uuid4()),
        'name': data.get('name', ''),
        'description': data.get('description', ''),
        'date': datetime.now().strftime('%Y-%m-%d'),
        'notes': [],
    }
    projects.append(new_project)
    save_projects(projects)
    return jsonify(new_project), 201

@app.route('/projects/<project_id>', methods=['PUT'])
def edit_project(project_id):
    data = request.json
    projects = load_projects()
    for p in projects:
        if p['id'] == project_id:
            p['name'] = data.get('name', p['name'])
            p['description'] = data.get('description', p['description'])
            save_projects(projects)
            return jsonify(p)
    return jsonify({'error': 'Project not found'}), 404

@app.route('/projects/<project_id>', methods=['DELETE'])
def delete_project(project_id):
    projects = load_projects()
    new_projects = [p for p in projects if p['id'] != project_id]
    if len(new_projects) == len(projects):
        return jsonify({'error': 'Project not found'}), 404
    save_projects(new_projects)
    return '', 204

@app.route('/projects/<project_id>/notes', methods=['GET'])
def get_project_notes(project_id):
    projects = load_projects()
    for p in projects:
        if p['id'] == project_id:
            return jsonify(p.get('notes', []))
    return jsonify({'error': 'Project not found'}), 404

@app.route('/projects/<project_id>/notes', methods=['POST'])
def add_project_note(project_id):
    data = request.json
    projects = load_projects()
    for p in projects:
        if p['id'] == project_id:
            note = {
                'id': str(uuid.uuid4()),
                'text': data.get('text', ''),
                'createdAt': datetime.now().isoformat(),
            }
            p.setdefault('notes', []).insert(0, note)
            save_projects(projects)
            return jsonify(note), 201
    return jsonify({'error': 'Project not found'}), 404

@app.route('/projects/<project_id>/notes/<note_id>', methods=['PUT'])
def edit_project_note(project_id, note_id):
    data = request.json
    projects = load_projects()
    for p in projects:
        if p['id'] == project_id:
            for n in p.get('notes', []):
                if n['id'] == note_id:
                    n['text'] = data.get('text', n['text'])
                    save_projects(projects)
                    return jsonify(n)
            return jsonify({'error': 'Note not found'}), 404
    return jsonify({'error': 'Project not found'}), 404

@app.route('/projects/<project_id>/notes/<note_id>', methods=['DELETE'])
def delete_project_note(project_id, note_id):
    projects = load_projects()
    for p in projects:
        if p['id'] == project_id:
            notes = p.get('notes', [])
            new_notes = [n for n in notes if n['id'] != note_id]
            if len(new_notes) == len(notes):
                return jsonify({'error': 'Note not found'}), 404
            p['notes'] = new_notes
            save_projects(projects)
            return '', 204
    return jsonify({'error': 'Project not found'}), 404

@app.route('/projects/<project_id>/chats', methods=['GET'])
def get_project_chats(project_id):
    projects = load_projects()
    for p in projects:
        if p['id'] == project_id:
            return jsonify(p.get('chats', []))
    return jsonify({'error': 'Project not found'}), 404

@app.route('/projects/<project_id>/chats', methods=['POST'])
def add_project_chat(project_id):
    data = request.json
    projects = load_projects()
    for p in projects:
        if p['id'] == project_id:
            chat = {
                'id': str(uuid.uuid4()),
                'title': data.get('title', f'Chat {len(p.get("chats", [])) + 1}'),
                'messages': data.get('messages', []),
                'createdAt': datetime.now().isoformat(),
            }
            if 'chats' not in p:
                p['chats'] = []
            p['chats'].insert(0, chat)
            save_projects(projects)
            return jsonify(chat), 201
    return jsonify({'error': 'Project not found'}), 404

@app.route('/projects/<project_id>/chats/<chat_id>', methods=['PUT'])
def edit_project_chat(project_id, chat_id):
    data = request.json
    projects = load_projects()
    for p in projects:
        if p['id'] == project_id:
            for c in p.get('chats', []):
                if c['id'] == chat_id:
                    c['title'] = data.get('title', c['title'])
                    if 'messages' in data:
                        c['messages'] = data['messages']
                    save_projects(projects)
                    return jsonify(c)
            return jsonify({'error': 'Chat not found'}), 404
    return jsonify({'error': 'Project not found'}), 404

FEEDBACK_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'backend', 'feedback.json')

@app.route('/feedback', methods=['POST'])
def submit_feedback():
    logger.info('POST /feedback endpoint hit (from /feedback route)')
    logger.info('POST /api/feedback endpoint hit (from /api/feedback route)')
    data = request.get_json()
    logger.info(f'Received feedback data: {data}')
    logger.info(f'FEEDBACK_FILE path: {FEEDBACK_FILE}')
    
    # Ensure all required fields are present with proper defaults
    feedback_entry = {
        'sessionId': data.get('sessionId'),
        'agentId': data.get('agentId'),
        'agentName': data.get('agentName', 'AI Assistant'),
        'answerId': data.get('answerId') or data.get('id'),
        'documentChunkIds': data.get('documentChunkIds'),
        'rating': data.get('rating'),
        'feedbackText': data.get('feedbackText') or data.get('reason', ''),
        'stars': data.get('stars') or data.get('rating'),
        'suggestion': data.get('suggestion', ''),
        'userId': data.get('userId'),
        'userEmail': data.get('userEmail', 'anonymous@user.com'),
        'timestamp': data.get('timestamp', datetime.utcnow().isoformat()),
        # Additional fields for dashboard
        'feedbackType': data.get('feedbackType', 'neutral'),
        'category': data.get('category', 'helpfulness'),
        'severity': data.get('severity', 'medium'),
        'modelUsed': data.get('modelUsed', 'Gemini 2.5 Flash'),
        'modelIcon': data.get('modelIcon', '🤖'),
        'responseTime': data.get('responseTime', 2.5),
        'sessionDuration': data.get('sessionDuration', 15),
        'answerText': data.get('answerText') or data.get('content', ''),
        'sessionTranscript': data.get('sessionTranscript', [])
    }
    
    # Log the complete feedback entry for debugging
    logger.info(f'Processed feedback entry: {feedback_entry}')
    try:
        # Ensure the backend directory exists
        os.makedirs(os.path.dirname(FEEDBACK_FILE), exist_ok=True)
        
        if not os.path.exists(FEEDBACK_FILE):
            logger.info('feedback.json does not exist, creating new file.')
            with open(FEEDBACK_FILE, 'w', encoding='utf-8') as f:
                json.dump([feedback_entry], f, indent=2, ensure_ascii=False)
        else:
            with open(FEEDBACK_FILE, 'r+', encoding='utf-8') as f:
                try:
                    feedbacks = json.load(f)
                    if not isinstance(feedbacks, list):
                        logger.warning('feedback.json does not contain a list, resetting to empty list')
                        feedbacks = []
                except Exception as e:
                    logger.error(f'Error loading existing feedbacks: {e}')
                    feedbacks = []
                
                feedbacks.append(feedback_entry)
                f.seek(0)
                json.dump(feedbacks, f, indent=2, ensure_ascii=False)
                f.truncate()
        
        logger.info(f'Feedback saved successfully. Total feedbacks: {len(feedbacks) if 'feedbacks' in locals() else 1}')
        return jsonify({'success': True, 'message': 'Feedback saved successfully'}), 200
    except Exception as e:
        logger.error(f"Error saving feedback: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/feedback', methods=['GET', 'POST'])
def get_feedback():
    if request.method == 'POST':
        return submit_feedback()
    
    """Get all feedback data for the dashboard"""
    logger.info(f'GET /api/feedback endpoint hit')
    logger.info(f'FEEDBACK_FILE path: {FEEDBACK_FILE}')
    
    try:
        if not os.path.exists(FEEDBACK_FILE):
            logger.info('feedback.json does not exist, returning empty list')
            return jsonify([])
        
        logger.info(f'Reading feedback file: {FEEDBACK_FILE}')
        with open(FEEDBACK_FILE, 'r', encoding='utf-8') as f:
            try:
                feedbacks = json.load(f)
                logger.info(f'Loaded {len(feedbacks) if isinstance(feedbacks, list) else 0} feedback entries')
                if not isinstance(feedbacks, list):
                    logger.warning('feedback.json does not contain a list, returning empty list')
                    feedbacks = []
            except Exception as e:
                logger.error(f'Error loading feedback data: {e}')
                feedbacks = []
        
        # Transform feedback data to match dashboard format
        transformed_feedbacks = []
        logger.info(f'Transforming {len(feedbacks)} feedback entries')
        for i, feedback in enumerate(feedbacks):
            logger.info(f'Processing feedback {i+1}: {feedback.get("feedbackText", "No text")[:50]}...')
            # Determine feedback type based on rating
            if feedback.get('rating', 0) >= 4:
                feedback_type = 'positive'
            elif feedback.get('rating', 0) <= 2:
                feedback_type = 'negative'
            else:
                feedback_type = 'neutral'
            
            # Determine category based on feedback text or default
            category = 'helpfulness'  # default
            if feedback.get('feedbackText'):
                text = feedback.get('feedbackText').lower()
                if any(word in text for word in ['speed', 'fast', 'slow', 'time']):
                    category = 'speed'
                elif any(word in text for word in ['accurate', 'correct', 'right', 'wrong']):
                    category = 'accuracy'
            
            # Determine severity based on rating
            if feedback.get('rating', 0) <= 2:
                severity = 'high'
            elif feedback.get('rating', 0) <= 3:
                severity = 'medium'
            else:
                severity = 'low'
            
            transformed_feedback = {
                'id': feedback.get('answerId') or feedback.get('id') or str(uuid.uuid4()),
                'session_id': feedback.get('sessionId'),
                'agent_id': feedback.get('agentId'),
                'agent_name': feedback.get('agentName', 'AI Assistant'),
                'rating': feedback.get('rating', 3),
                'category': feedback.get('category', category),
                'severity': feedback.get('severity', severity),
                'feedback_text': feedback.get('feedbackText') or feedback.get('reason', ''),
                'timestamp': feedback.get('timestamp'),
                'user_email': feedback.get('userEmail', feedback.get('userId', 'anonymous@user.com')),
                'response_time': feedback.get('responseTime', 2.5),
                'session_duration': feedback.get('sessionDuration', 15),
                'feedback_type': feedback.get('feedbackType', feedback_type),
                'message_id': feedback.get('answerId') or feedback.get('id'),
                'model_used': feedback.get('modelUsed', 'Gemini 2.5 Flash'),
                'model_icon': feedback.get('modelIcon', '🤖'),
                'answer_text': feedback.get('answerText') or feedback.get('content', '')
            }
            transformed_feedbacks.append(transformed_feedback)
        
        logger.info(f'Returning {len(transformed_feedbacks)} transformed feedback entries')
        return jsonify(transformed_feedbacks)
    except Exception as e:
        logger.error(f"Error loading feedback data: {e}")
        return jsonify({'error': str(e)}), 500

# Feedback deletion is disabled to ensure data persistence
# @app.route('/feedback', methods=['DELETE'])
# def delete_feedback():
#     """Delete feedback by answerId - DISABLED for data persistence"""
#     return jsonify({'error': 'Feedback deletion is disabled to ensure data persistence'}), 403

@app.errorhandler(404)
def not_found_error(error):
    return jsonify({'error': 'Not found', 'success': False}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error', 'success': False}), 500

# Monitoring Dashboard API Endpoints
@app.route('/api/monitoring', methods=['GET', 'POST'])
def get_monitoring_data():
    if request.method == 'POST':
        """Log monitoring data from frontend"""
        try:
            data = request.json
            user_email = data.get('userEmail', 'anonymous@user.com')
            action = data.get('action', 'unknown')
            details = data.get('details', {})
            
            log_user_activity(
                user_email=user_email,
                action=action,
                details=details
            )
            
            return jsonify({'message': 'Activity logged successfully'}), 200
        except Exception as e:
            logger.error(f"Error logging activity: {e}")
            return jsonify({'error': 'Failed to log activity'}), 500
    
    """Get monitoring data for dashboard"""
    try:
        # Get query parameters for filtering
        date_from = request.args.get('date_from')
        date_to = request.args.get('date_to')
        user_email = request.args.get('user_email')
        model_name = request.args.get('model_name')
        agent_id = request.args.get('agent_id')
        
        # Filter data based on parameters
        filtered_data = {
            'model_usage': monitoring_data['model_usage'].copy(),
            'user_activity': monitoring_data['user_activity'].copy(),
            'session_metrics': monitoring_data['session_metrics'].copy(),
            'agent_usage': monitoring_data['agent_usage'].copy(),
            'system_metrics': monitoring_data['system_metrics'].copy()
        }
        
        # Apply filters
        if date_from:
            filtered_data['model_usage'] = [entry for entry in filtered_data['model_usage'] if entry['timestamp'] >= date_from]
            filtered_data['user_activity'] = [entry for entry in filtered_data['user_activity'] if entry['timestamp'] >= date_from]
            filtered_data['session_metrics'] = [entry for entry in filtered_data['session_metrics'] if entry['timestamp'] >= date_from]
            filtered_data['agent_usage'] = [entry for entry in filtered_data['agent_usage'] if entry['timestamp'] >= date_from]
            filtered_data['system_metrics'] = [entry for entry in filtered_data['system_metrics'] if entry['timestamp'] >= date_from]
        
        if date_to:
            filtered_data['model_usage'] = [entry for entry in filtered_data['model_usage'] if entry['timestamp'] <= date_to]
            filtered_data['user_activity'] = [entry for entry in filtered_data['user_activity'] if entry['timestamp'] <= date_to]
            filtered_data['session_metrics'] = [entry for entry in filtered_data['session_metrics'] if entry['timestamp'] <= date_to]
            filtered_data['agent_usage'] = [entry for entry in filtered_data['agent_usage'] if entry['timestamp'] <= date_to]
            filtered_data['system_metrics'] = [entry for entry in filtered_data['system_metrics'] if entry['timestamp'] <= date_to]
        
        if user_email:
            filtered_data['model_usage'] = [entry for entry in filtered_data['model_usage'] if entry['user_email'] == user_email]
            filtered_data['user_activity'] = [entry for entry in filtered_data['user_activity'] if entry['user_email'] == user_email]
            filtered_data['session_metrics'] = [entry for entry in filtered_data['session_metrics'] if entry['user_email'] == user_email]
            filtered_data['agent_usage'] = [entry for entry in filtered_data['agent_usage'] if entry['user_email'] == user_email]
        
        if model_name:
            filtered_data['model_usage'] = [entry for entry in filtered_data['model_usage'] if entry['model_name'] == model_name]
            filtered_data['session_metrics'] = [entry for entry in filtered_data['session_metrics'] if entry['model_used'] == model_name]
        
        if agent_id:
            filtered_data['model_usage'] = [entry for entry in filtered_data['model_usage'] if entry['agent_id'] == agent_id]
            filtered_data['user_activity'] = [entry for entry in filtered_data['user_activity'] if entry['agent_id'] == agent_id]
            filtered_data['session_metrics'] = [entry for entry in filtered_data['session_metrics'] if entry['agent_id'] == agent_id]
            filtered_data['agent_usage'] = [entry for entry in filtered_data['agent_usage'] if entry['agent_id'] == agent_id]
        
        # Debug: Print session IDs in monitoring data
        session_ids_in_monitoring = [entry['session_id'] for entry in filtered_data['session_metrics'] if 'session_id' in entry]
        print(f"DEBUG: Session IDs in monitoring data: {session_ids_in_monitoring}")
        print(f"DEBUG: Available sessions: {list(sessions.keys())}")
        
        return jsonify(filtered_data)
        
    except Exception as e:
        logger.error(f"Error getting monitoring data: {e}")
        return jsonify({'error': 'Failed to get monitoring data'}), 500

@app.route('/api/monitoring/summary', methods=['GET'])
def get_monitoring_summary():
    """Get monitoring summary statistics"""
    try:
        # Calculate summary statistics
        total_model_usage = len(monitoring_data['model_usage'])
        total_user_activity = len(monitoring_data['user_activity'])
        total_sessions = len(monitoring_data['session_metrics'])
        total_agent_usage = len(monitoring_data['agent_usage'])
        
        # Model usage statistics
        model_counts = {}
        for entry in monitoring_data['model_usage']:
            model = entry['model_name']
            model_counts[model] = model_counts.get(model, 0) + 1
        
        # User activity statistics
        user_counts = {}
        for entry in monitoring_data['user_activity']:
            user = entry['user_email']
            user_counts[user] = user_counts.get(user, 0) + 1
        
        # Agent usage statistics
        agent_counts = {}
        for entry in monitoring_data['agent_usage']:
            agent = entry['agent_name']
            agent_counts[agent] = agent_counts.get(agent, 0) + 1
        
        # Response time statistics
        response_times = [entry['response_time'] for entry in monitoring_data['model_usage'] if entry.get('response_time')]
        avg_response_time = sum(response_times) / len(response_times) if response_times else 0
        
        # Success rate
        successful_requests = len([entry for entry in monitoring_data['model_usage'] if entry.get('success', True)])
        success_rate = (successful_requests / total_model_usage * 100) if total_model_usage > 0 else 0
        
        summary = {
            'total_model_usage': total_model_usage,
            'total_user_activity': total_user_activity,
            'total_sessions': total_sessions,
            'total_agent_usage': total_agent_usage,
            'model_counts': model_counts,
            'user_counts': user_counts,
            'agent_counts': agent_counts,
            'avg_response_time': round(avg_response_time, 2),
            'success_rate': round(success_rate, 2)
        }
        
        return jsonify(summary)
        
    except Exception as e:
        logger.error(f"Error getting monitoring summary: {e}")
        return jsonify({'error': 'Failed to get monitoring summary'}), 500

@app.route('/api/monitoring/clear', methods=['POST'])
def clear_monitoring_data():
    """Clear all monitoring data (admin only)"""
    try:
        global monitoring_data
        monitoring_data = {
            'model_usage': [],
            'user_activity': [],
            'session_metrics': [],
            'agent_usage': [],
            'system_metrics': []
        }
        save_monitoring_data()
        return jsonify({'message': 'Monitoring data cleared successfully'}), 200
    except Exception as e:
        logger.error(f"Error clearing monitoring data: {e}")
        return jsonify({'error': 'Failed to clear monitoring data'}), 500

@app.route('/api/monitoring/agent/<agent_id>', methods=['GET'])
def get_agent_details(agent_id):
    """Get detailed information about a specific agent"""
    try:
        if agent_id in AGENTS_DATA:
            agent_info = AGENTS_DATA[agent_id].copy()
            
            # Get usage statistics for this agent
            agent_usage = [entry for entry in monitoring_data['agent_usage'] if entry['agent_id'] == agent_id]
            model_usage = [entry for entry in monitoring_data['model_usage'] if entry['agent_id'] == agent_id]
            user_activity = [entry for entry in monitoring_data['user_activity'] if entry['agent_id'] == agent_id]
            
            # Calculate statistics
            total_interactions = len(agent_usage)
            unique_users = len(set(entry['user_email'] for entry in agent_usage))
            total_sessions = len(set(entry['session_id'] for entry in agent_usage))
            
            # Get recent sessions
            recent_sessions = []
            for entry in agent_usage[:10]:  # Last 10 interactions
                session_id = entry['session_id']
                if session_id in sessions:
                    session_info = sessions[session_id]
                    recent_sessions.append({
                        'session_id': session_id,
                        'title': session_info.get('title', 'Untitled'),
                        'created_at': session_info.get('createdAt'),
                        'message_count': len(session_info.get('messages', [])),
                        'user_email': entry['user_email']
                    })
            
            return jsonify({
                'agent_info': agent_info,
                'statistics': {
                    'total_interactions': total_interactions,
                    'unique_users': unique_users,
                    'total_sessions': total_sessions,
                    'model_usage_count': len(model_usage),
                    'user_activity_count': len(user_activity)
                },
                'recent_sessions': recent_sessions
            })
        else:
            return jsonify({'error': 'Agent not found'}), 404
    except Exception as e:
        logger.error(f"Error getting agent details: {e}")
        return jsonify({'error': 'Failed to get agent details'}), 500

@app.route('/api/monitoring/session/<session_id>', methods=['GET'])
def get_session_details(session_id):
    """Get detailed information about a specific session"""
    try:
        print(f"DEBUG: Looking for session_id: {session_id}")
        print(f"DEBUG: Available session IDs: {list(sessions.keys())}")
        print(f"DEBUG: Session ID in sessions: {session_id in sessions}")
        
        if session_id in sessions:
            session_info = sessions[session_id]
            
            # Get monitoring data for this session
            session_metrics = [entry for entry in monitoring_data['session_metrics'] if entry['session_id'] == session_id]
            model_usage = [entry for entry in monitoring_data['model_usage'] if entry['session_id'] == session_id]
            user_activity = [entry for entry in monitoring_data['user_activity'] if entry['session_id'] == session_id]
            agent_usage = [entry for entry in monitoring_data['agent_usage'] if entry['session_id'] == session_id]
            
            # Calculate session statistics
            total_messages = len(session_info.get('messages', []))
            created_at = datetime.fromisoformat(session_info.get('createdAt', datetime.now().isoformat()))
            session_duration = (datetime.now() - created_at).total_seconds()
            
            # Get message statistics
            user_messages = [msg for msg in session_info.get('messages', []) if msg.get('role') == 'user']
            assistant_messages = [msg for msg in session_info.get('messages', []) if msg.get('role') == 'assistant']
            
            return jsonify({
                'session_info': session_info,
                'statistics': {
                    'total_messages': total_messages,
                    'user_messages': len(user_messages),
                    'assistant_messages': len(assistant_messages),
                    'session_duration_seconds': session_duration,
                    'created_at': session_info.get('createdAt'),
                    'model_usage_count': len(model_usage),
                    'user_activity_count': len(user_activity),
                    'agent_usage_count': len(agent_usage)
                },
                'monitoring_data': {
                    'session_metrics': session_metrics,
                    'model_usage': model_usage,
                    'user_activity': user_activity,
                    'agent_usage': agent_usage
                }
            })
        else:
            return jsonify({'error': 'Session not found'}), 404
    except Exception as e:
        logger.error(f"Error getting session details: {e}")
        return jsonify({'error': 'Failed to get session details'}), 500

@app.route('/upload', methods=['POST'])
def upload_any():
    """Handle uploads of any file type, save them, and return filenames."""
    if request.method == 'OPTIONS':
        return '', 200
    try:
        if 'files' not in request.files:
            logger.error("No files part in request")
            return jsonify({'error': 'No files part', 'success': False}), 400
        files = request.files.getlist('files')
        if not files or all(f.filename == '' for f in files):
            logger.error("No selected files")
            return jsonify({'error': 'No selected files', 'success': False}), 400
        saved_filenames = []
        for file in files:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            logger.info(f"Saving file: {filepath}")
            file.save(filepath)
            saved_filenames.append(filename)
        return jsonify({
            'message': 'Files uploaded successfully',
            'results': [{'filename': fn} for fn in saved_filenames],
            'success': True
        })
    except Exception as e:
        logger.error(f"Unexpected error in upload_any route: {str(e)}")
        return jsonify({'error': f'Unexpected error: {str(e)}', 'success': False}), 500

@app.route('/whisper/stt', methods=['POST', 'OPTIONS'])
def whisper_stt():
    if request.method == 'OPTIONS':
        return '', 200
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400

        audio_file = request.files['audio']
        import os, uuid

        # Ensure the temp audio directory exists
        temp_audio_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'backend', 'audio_temp')
        os.makedirs(temp_audio_dir, exist_ok=True)

        filename = audio_file.filename
        ext = os.path.splitext(filename)[1] or '.webm'
        temp_path = os.path.join(temp_audio_dir, f"{uuid.uuid4().hex}{ext}")

        audio_file.save(temp_path)

        if not WHISPER_AVAILABLE:
            os.remove(temp_path)
            return jsonify({'error': 'Whisper is not available for audio transcription.'}), 500
        try:
            result = whisper_model.transcribe(temp_path, language=None)  # auto-detect language
            transcript = result['text']
            detected_language = result.get('language', 'unknown')
        except Exception as e:
            logger.error(f"Whisper transcription failed for uploaded audio: {e}")
            os.remove(temp_path)
            return jsonify({'error': str(e)}), 500

        os.remove(temp_path)
        response = jsonify({'transcript': transcript, 'language': detected_language})
        response.headers['Access-Control-Allow-Origin'] = 'http://localhost:3000'
        response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        response.headers['Access-Control-Allow-Credentials'] = 'true'
        return response

    except Exception as e:
        logger.error(f"Whisper STT error: {e}")
        return jsonify({'error': str(e)}), 500

def extract_meaningful_highlight(chunk, question, window=3):
    import re
    sentences = re.split(r'(?<=[.!?]) +', chunk)
    question_words = set(re.findall(r'\w+', question.lower()))
    scores = []
    for i, sent in enumerate(sentences):
        sent_words = set(re.findall(r'\w+', sent.lower()))
        overlap = len(question_words & sent_words)
        scores.append((overlap, i))
    best_score, best_idx = max(scores) if scores else (0, 0)
    start = max(0, best_idx - window//2)
    end = min(len(sentences), best_idx + window//2 + 1)
    highlight = ' '.join(sentences[start:end]).strip()
    # Filter: skip highlights that are only a single character or match (a), (b), (c), etc.
    if re.fullmatch(r'\([a-zA-Z]\)', highlight) or len(highlight.split()) < 2:
        # Try to expand window if possible
        start = max(0, best_idx - window)
        end = min(len(sentences), best_idx + window + 1)
        highlight = ' '.join(sentences[start:end]).strip()
        # If still not meaningful, fallback to first 200 chars
        if re.fullmatch(r'\([a-zA-Z]\)', highlight) or len(highlight.split()) < 2:
            return chunk[:200]
    return highlight if highlight else chunk[:200]

def load_documents_from_directory(directory_path: str) -> list[Document]:
    """Load all supported documents from a directory and convert them to LangChain Documents"""
    from langchain_core.documents import Document
    import docx
    import csv
    import mimetypes
    supported_docs = []
    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        if filename.lower().endswith('.pdf'):
            try:
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                doc = Document(page_content=text, metadata={"source": filename})
                supported_docs.append(doc)
            except Exception as e:
                logger.error(f"Error loading PDF {filename}: {str(e)}")
        elif filename.lower().endswith('.txt'):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                doc = Document(page_content=text, metadata={"source": filename})
                supported_docs.append(doc)
            except Exception as e:
                logger.error(f"Error loading TXT {filename}: {str(e)}")
        elif filename.lower().endswith('.docx'):
            try:
                docx_file = docx.Document(file_path)
                text = "\n".join([p.text for p in docx_file.paragraphs])
                doc = Document(page_content=text, metadata={"source": filename})
                supported_docs.append(doc)
            except Exception as e:
                logger.error(f"Error loading DOCX {filename}: {str(e)}")
        elif filename.lower().endswith('.csv'):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    text = "\n".join([", ".join(row) for row in reader])
                doc = Document(page_content=text, metadata={"source": filename})
                supported_docs.append(doc)
            except Exception as e:
                logger.error(f"Error loading CSV {filename}: {str(e)}")
        elif filename.lower().endswith(('.png', '.jpg', '.jpeg')):
            try:
                from backend.pdf_utils import convert_image_to_text
                with open(file_path, 'rb') as imgf:
                    image_data = imgf.read()
                text = convert_image_to_text(image_data)
                doc = Document(page_content=text, metadata={"source": filename})
                supported_docs.append(doc)
            except Exception as e:
                logger.error(f"Error loading image {filename}: {str(e)}")
        elif filename.lower().endswith('.mp3'):
            try:
                # Placeholder: add your audio transcription logic here
                text = f"[Audio file: {filename}]"
                doc = Document(page_content=text, metadata={"source": filename})
                supported_docs.append(doc)
            except Exception as e:
                logger.error(f"Error loading audio {filename}: {str(e)}")
        else:
            logger.info(f"Skipping unsupported file type: {filename}")
    return supported_docs

def load_txt_as_documents(txt_filename):
    txt_dir = app.config['UPLOAD_FOLDER']
    txt_path = os.path.join(txt_dir, txt_filename)
    docs = []
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            text = f.read()
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        chunks = splitter.split_text(text)
        docs = [Document(page_content=chunk, metadata={"source": txt_filename, "file_type": "txt"}) for chunk in chunks]
        logger.info(f"Loaded {len(docs)} chunks from TXT {txt_filename}")
    except Exception as e:
        logger.warning(f"Failed to load TXT {txt_filename}: {e}")
    return docs

def load_docx_as_documents(docx_filename):
    import docx2txt
    docx_dir = app.config['UPLOAD_FOLDER']
    docx_path = os.path.join(docx_dir, docx_filename)
    docs = []
    try:
        text = docx2txt.process(docx_path)
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        chunks = splitter.split_text(text)
        docs = [Document(page_content=chunk, metadata={"source": docx_filename, "file_type": "docx"}) for chunk in chunks]
        logger.info(f"Loaded {len(docs)} chunks from DOCX {docx_filename}")
    except Exception as e:
        logger.warning(f"Failed to load DOCX {docx_filename}: {e}")
    return docs

def load_csv_as_documents(csv_filename):
    import csv
    csv_dir = app.config['UPLOAD_FOLDER']
    csv_path = os.path.join(csv_dir, csv_filename)
    docs = []
    try:
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            rows = list(reader)
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        for i, row in enumerate(rows):
            row_text = ', '.join(row)
            splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            chunks = splitter.split_text(row_text)
            docs.extend([Document(page_content=chunk, metadata={"source": csv_filename, "file_type": "csv", "row": i}) for chunk in chunks])
        logger.info(f"Loaded {len(docs)} chunks from CSV {csv_filename}")
    except Exception as e:
        logger.warning(f"Failed to load CSV {csv_filename}: {e}")
    return docs

def load_image_as_documents(image_filename):
    from backend.pdf_utils import convert_image_to_text
    img_dir = app.config['UPLOAD_FOLDER']
    img_path = os.path.join(img_dir, image_filename)
    docs = []
    try:
        with open(img_path, 'rb') as imgf:
            image_data = imgf.read()
        text = convert_image_to_text(image_data)
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        chunks = splitter.split_text(text)
        docs = [Document(page_content=chunk, metadata={"source": image_filename, "file_type": "image"}) for chunk in chunks]
        logger.info(f"Loaded {len(docs)} chunks from image {image_filename}")
    except Exception as e:
        logger.warning(f"Failed to load image {image_filename}: {e}")
    return docs

def load_audio_as_documents(audio_filename):
    audio_dir = app.config['UPLOAD_FOLDER']
    audio_path = os.path.join(audio_dir, audio_filename)
    docs = []
    try:
        transcript_path = os.path.splitext(audio_path)[0] + '.txt'
        if os.path.exists(transcript_path):
            with open(transcript_path, 'r', encoding='utf-8') as tf:
                text = tf.read()
        else:
            text = extract_audio_transcript(audio_path)
            if text:
                with open(transcript_path, 'w', encoding='utf-8') as tf:
                    tf.write(text)
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        chunks = splitter.split_text(text)
        docs = [Document(page_content=chunk, metadata={"source": audio_filename, "file_type": "audio"}) for chunk in chunks]
        logger.info(f"Loaded {len(docs)} chunks from audio {audio_filename}")
    except Exception as e:
        logger.warning(f"Failed to load audio {audio_filename}: {e}")
    return docs

# Update load_documents_from_directory to use these loaders for each file type

def load_documents_from_directory(directory_path: str) -> list[Document]:
    """Load all supported documents from a directory and convert them to LangChain Documents"""
    supported_docs = []
    for filename in os.listdir(directory_path):
        if filename.lower().endswith('.pdf'):
            supported_docs.extend(load_pdf_as_documents(filename))
        elif filename.lower().endswith('.txt'):
            supported_docs.extend(load_txt_as_documents(filename))
        elif filename.lower().endswith('.docx'):
            supported_docs.extend(load_docx_as_documents(filename))
        elif filename.lower().endswith('.csv'):
            supported_docs.extend(load_csv_as_documents(filename))
        elif filename.lower().endswith(('.png', '.jpg', '.jpeg')):
            supported_docs.extend(load_image_as_documents(filename))
        elif filename.lower().endswith('.mp3'):
            supported_docs.extend(load_audio_as_documents(filename))
        else:
            logger.info(f"Skipping unsupported file type: {filename}")
    return supported_docs

# Monitoring and Analytics
MONITORING_FILE = 'monitoring_data.json'

# Initialize monitoring data structure
monitoring_data = {
    'model_usage': [],
    'user_activity': [],
    'session_metrics': [],
    'agent_usage': [],
    'system_metrics': []
}

def save_monitoring_data():
    """Save monitoring data to JSON file."""
    try:
        with open(MONITORING_FILE, 'w') as f:
            json.dump(monitoring_data, f, indent=4)
        logger.info("Monitoring data saved to file.")
    except Exception as e:
        logger.error(f"Error saving monitoring data: {e}")

def load_monitoring_data():
    """Load monitoring data from JSON file."""
    global monitoring_data
    if os.path.exists(MONITORING_FILE):
        try:
            with open(MONITORING_FILE, 'r') as f:
                monitoring_data = json.load(f)
            logger.info("Monitoring data loaded from file.")
        except Exception as e:
            logger.error(f"Error loading monitoring data: {e}")

def log_model_usage(user_email, model_name, session_id, agent_id=None, response_time=None, tokens_used=None, success=True, error_message=None):
    """Log model usage for monitoring."""
    usage_entry = {
        'id': str(uuid.uuid4()),
        'user_email': user_email,
        'model_name': model_name,
        'session_id': session_id,
        'agent_id': agent_id,
        'response_time': response_time,
        'tokens_used': tokens_used,
        'success': success,
        'error_message': error_message,
        'timestamp': datetime.now().isoformat()
    }
    monitoring_data['model_usage'].append(usage_entry)
    save_monitoring_data()

def log_user_activity(user_email, action, session_id=None, agent_id=None, details=None):
    """Log user activity for monitoring."""
    activity_entry = {
        'id': str(uuid.uuid4()),
        'user_email': user_email,
        'action': action,  # 'login', 'logout', 'send_message', 'create_session', 'delete_session', etc.
        'session_id': session_id,
        'agent_id': agent_id,
        'details': details,
        'timestamp': datetime.now().isoformat()
    }
    monitoring_data['user_activity'].append(activity_entry)
    save_monitoring_data()

def log_session_metrics(session_id, user_email, agent_id, message_count, session_duration, model_used):
    """Log session metrics for monitoring."""
    # Get agent name if agent_id is provided
    agent_name = None
    if agent_id and agent_id in AGENTS_DATA:
        agent_name = AGENTS_DATA[agent_id].get('name', 'Unknown Agent')
    
    metrics_entry = {
        'id': str(uuid.uuid4()),
        'session_id': session_id,
        'user_email': user_email,
        'agent_id': agent_id,
        'agent_name': agent_name,
        'message_count': message_count,
        'session_duration': session_duration,  # in seconds
        'model_used': model_used,
        'timestamp': datetime.now().isoformat()
    }
    monitoring_data['session_metrics'].append(metrics_entry)
    save_monitoring_data()

def log_agent_usage(agent_id, agent_name, user_email, session_id, usage_type='interaction'):
    """Log agent usage for monitoring."""
    usage_entry = {
        'id': str(uuid.uuid4()),
        'agent_id': agent_id,
        'agent_name': agent_name,
        'user_email': user_email,
        'session_id': session_id,
        'usage_type': usage_type,  # 'interaction', 'creation', 'modification'
        'timestamp': datetime.now().isoformat()
    }
    monitoring_data['agent_usage'].append(usage_entry)
    save_monitoring_data()

def log_system_metrics(metric_type, value, details=None):
    """Log system metrics for monitoring."""
    metrics_entry = {
        'id': str(uuid.uuid4()),
        'metric_type': metric_type,  # 'memory_usage', 'cpu_usage', 'disk_usage', 'api_calls', etc.
        'value': value,
        'details': details,
        'timestamp': datetime.now().isoformat()
    }
    monitoring_data['system_metrics'].append(metrics_entry)
    save_monitoring_data()

# Load monitoring data on startup
load_monitoring_data()

# GitHub OAuth Configuration
GITHUB_CLIENT_ID = os.getenv('GITHUB_CLIENT_ID', 'YOUR_GITHUB_CLIENT_ID_HERE')
GITHUB_CLIENT_SECRET = os.getenv('GITHUB_CLIENT_SECRET', 'YOUR_GITHUB_CLIENT_SECRET_HERE')

@app.route('/api/github/auth', methods=['POST'])
def github_auth():
    """Handle GitHub OAuth callback and exchange code for access token"""
    try:
        data = request.get_json()
        code = data.get('code')
        
        if not code:
            return jsonify({'error': 'Authorization code is required'}), 400
        
        # Exchange code for access token
        token_response = requests.post('https://github.com/login/oauth/access_token', {
            'client_id': GITHUB_CLIENT_ID,
            'client_secret': GITHUB_CLIENT_SECRET,
            'code': code
        }, headers={'Accept': 'application/json'})
        
        if token_response.status_code != 200:
            return jsonify({'error': 'Failed to exchange code for token'}), 400
        
        token_data = token_response.json()
        access_token = token_data.get('access_token')
        
        if not access_token:
            return jsonify({'error': 'Failed to get access token'}), 400
        
        # Get user information
        user_response = requests.get('https://api.github.com/user', {
            'access_token': access_token
        }, headers={'Authorization': f'token {access_token}'})
        
        if user_response.status_code != 200:
            return jsonify({'error': 'Failed to get user information'}), 400
        
        user_data = user_response.json()
        
        # Get user email
        emails_response = requests.get('https://api.github.com/user/emails', {
            'access_token': access_token
        }, headers={'Authorization': f'token {access_token}'})
        
        emails_data = emails_response.json() if emails_response.status_code == 200 else []
        primary_email = next((email['email'] for email in emails_data if email['primary']), user_data.get('email'))
        
        user_info = {
            'id': user_data['id'],
            'login': user_data['login'],
            'name': user_data.get('name') or user_data['login'],
            'email': primary_email,
            'avatar_url': user_data['avatar_url'],
            'access_token': access_token
        }
        
        return jsonify({
            'access_token': access_token,
            'user': user_info
        })
        
    except Exception as e:
        logger.error(f"GitHub OAuth error: {e}")
        return jsonify({'error': 'GitHub authentication failed'}), 500

# OTP Functions
def generate_otp():
    """Generate a 6-digit OTP"""
    return ''.join([str(secrets.randbelow(10)) for _ in range(6)])

def send_otp_email(email, otp, email_type='password_reset'):
    """Send OTP via email"""
    try:
        if email_type == 'account_creation':
            subject = 'Account Creation OTP - Unified Knowledge Platform'
            body_text = f'''
Hello,

You have requested to create a new account on Unified Knowledge Platform.

Your OTP (One-Time Password) is: {otp}

This OTP is valid for 10 minutes. Please do not share this OTP with anyone.

If you did not request to create this account, please ignore this email.

Best regards,
Unified Knowledge Platform Team
            '''
            html_content = f'''
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
        .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
        .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; text-align: center; border-radius: 10px 10px 0 0; }}
        .content {{ background: #f9f9f9; padding: 20px; border-radius: 0 0 10px 10px; }}
        .otp {{ background: #667eea; color: white; padding: 15px; text-align: center; font-size: 24px; font-weight: bold; border-radius: 5px; margin: 20px 0; }}
        .warning {{ background: #fff3cd; border: 1px solid #ffeaa7; padding: 10px; border-radius: 5px; margin: 20px 0; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>Account Creation OTP</h1>
            <p>Unified Knowledge Platform</p>
        </div>
        <div class="content">
            <p>Hello,</p>
            <p>You have requested to create a new account on Unified Knowledge Platform.</p>
            <div class="otp">Your OTP: {otp}</div>
            <p>This OTP is valid for <strong>10 minutes</strong>. Please do not share this OTP with anyone.</p>
            <div class="warning">
                <strong>Security Notice:</strong> If you did not request to create this account, please ignore this email.
            </div>
            <p>Best regards,<br>Unified Knowledge Platform Team</p>
        </div>
    </div>
</body>
</html>
            '''
        elif email_type == 'account_deletion':
            subject = 'Account Deletion OTP - Unified Knowledge Platform'
            body_text = f'''
Hello,

We're sorry to see you go! You have requested to delete your Unified Knowledge Platform account.

Your OTP (One-Time Password) is: {otp}

This OTP is valid for 10 minutes. Please do not share this OTP with anyone.

⚠️ WARNING: This action is irreversible and will permanently delete your account and all associated data including:
• Your profile information
• All your projects and notes
• Chat history and conversations
• Uploaded documents and files
• Account settings and preferences

If you did not request to delete your account, please ignore this email and ensure your account is secure.

If you change your mind, you can always create a new account in the future.

We hope to see you again!

Best regards,
Unified Knowledge Platform Team
            '''
            html_content = f'''
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
        .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
        .header {{ background: linear-gradient(135deg, #dc3545 0%, #c82333 100%); color: white; padding: 20px; text-align: center; border-radius: 10px 10px 0 0; }}
        .content {{ background: #f9f9f9; padding: 20px; border-radius: 0 0 10px 10px; }}
        .otp {{ background: #dc3545; color: white; padding: 15px; text-align: center; font-size: 24px; font-weight: bold; border-radius: 5px; margin: 20px 0; }}
        .danger {{ background: #f8d7da; border: 1px solid #f5c6cb; padding: 15px; border-radius: 5px; margin: 20px 0; color: #721c24; }}
        .warning {{ background: #fff3cd; border: 1px solid #ffeaa7; padding: 10px; border-radius: 5px; margin: 20px 0; }}
        .goodbye {{ background: #e8f5e8; border: 1px solid #c3e6c3; padding: 15px; border-radius: 5px; margin: 20px 0; color: #2d5a2d; }}
        .data-list {{ margin: 10px 0; padding-left: 20px; }}
        .data-list li {{ margin: 5px 0; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>Account Deletion OTP</h1>
            <p>Unified Knowledge Platform</p>
        </div>
        <div class="content">
            <p>Hello,</p>
            <p><strong>We're sorry to see you go!</strong> You have requested to delete your Unified Knowledge Platform account.</p>
            <div class="otp">Your OTP: {otp}</div>
            <p>This OTP is valid for <strong>10 minutes</strong>. Please do not share this OTP with anyone.</p>
            <div class="danger">
                <strong>⚠️ WARNING:</strong> This action is irreversible and will permanently delete your account and all associated data including:
                <ul class="data-list">
                    <li>Your profile information</li>
                    <li>All your projects and notes</li>
                    <li>Chat history and conversations</li>
                    <li>Uploaded documents and files</li>
                    <li>Account settings and preferences</li>
                </ul>
            </div>
            <div class="warning">
                <strong>Security Notice:</strong> If you did not request to delete your account, please ignore this email and ensure your account is secure.
            </div>
            <div class="goodbye">
                <strong>💝 Farewell Message:</strong> If you change your mind, you can always create a new account in the future. We hope to see you again!
            </div>
            <p>Best regards,<br>Unified Knowledge Platform Team</p>
        </div>
    </div>
</body>
</html>
            '''
        else:  # password_reset
            subject = 'Password Reset OTP - Unified Knowledge Platform'
            body_text = f'''
Hello,

You have requested a password reset for your Unified Knowledge Platform account.

Your OTP (One-Time Password) is: {otp}

This OTP is valid for 10 minutes. Please do not share this OTP with anyone.

If you did not request this password reset, please ignore this email.

Best regards,
Unified Knowledge Platform Team
            '''
            html_content = f'''
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
        .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
        .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; text-align: center; border-radius: 10px 10px 0 0; }}
        .content {{ background: #f9f9f9; padding: 20px; border-radius: 0 0 10px 10px; }}
        .otp {{ background: #667eea; color: white; padding: 15px; text-align: center; font-size: 24px; font-weight: bold; border-radius: 5px; margin: 20px 0; }}
        .warning {{ background: #fff3cd; border: 1px solid #ffeaa7; padding: 10px; border-radius: 5px; margin: 20px 0; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>Password Reset OTP</h1>
            <p>Unified Knowledge Platform</p>
        </div>
        <div class="content">
            <p>Hello,</p>
            <p>You have requested a password reset for your Unified Knowledge Platform account.</p>
            <div class="otp">Your OTP: {otp}</div>
            <p>This OTP is valid for <strong>10 minutes</strong>. Please do not share this OTP with anyone.</p>
            <div class="warning">
                <strong>Security Notice:</strong> If you did not request this password reset, please ignore this email and ensure your account is secure.
            </div>
            <p>Best regards,<br>Unified Knowledge Platform Team</p>
        </div>
    </div>
</body>
</html>
            '''
        
        msg = Message(
            subject=subject,
            recipients=[email],
            body=body_text,
            html=html_content
        )
        mail.send(msg)
        return True
    except Exception as e:
        logger.error(f"Failed to send OTP email: {e}")
        return False

@app.route('/api/otp/send', methods=['POST'])
def send_otp():
    """Send OTP to user's email"""
    try:
        data = request.get_json()
        email = data.get('email')
        email_type = data.get('type', 'password_reset')  # Default to password_reset
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
        
        # Check if user exists (you can modify this based on your user storage)
        # For now, we'll assume the user exists if they provide an email
        
        # Generate OTP
        otp = generate_otp()
        expiry_time = datetime.now() + timedelta(minutes=10)
        
        # Store OTP with expiry and type
        otp_storage[email] = {
            'otp': otp,
            'expiry': expiry_time,
            'attempts': 0,
            'type': email_type
        }
        
        # Send OTP via email with the correct type
        if send_otp_email(email, otp, email_type):
            return jsonify({
                'message': 'OTP sent successfully',
                'email': email
            })
        else:
            return jsonify({'error': 'Failed to send OTP email'}), 500
            
    except Exception as e:
        logger.error(f"OTP send error: {e}")
        return jsonify({'error': 'Failed to send OTP'}), 500

@app.route('/api/otp/verify', methods=['POST'])
def verify_otp():
    """Verify OTP and allow password reset"""
    try:
        data = request.get_json()
        email = data.get('email')
        otp = data.get('otp')
        otp_type = data.get('type', 'password_reset')  # Default to password_reset
        
        if not email or not otp:
            return jsonify({'error': 'Email and OTP are required'}), 400
        
        # Check if OTP exists and is not expired
        if email not in otp_storage:
            return jsonify({'error': 'OTP not found or expired'}), 400
        
        otp_data = otp_storage[email]
        
        # Check if OTP type matches
        if otp_data.get('type') != otp_type:
            return jsonify({'error': 'Invalid OTP type'}), 400
        
        # Check if OTP is expired
        if datetime.now() > otp_data['expiry']:
            del otp_storage[email]
            return jsonify({'error': 'OTP has expired'}), 400
        
        # Check if too many attempts
        if otp_data['attempts'] >= 3:
            del otp_storage[email]
            return jsonify({'error': 'Too many failed attempts. Please request a new OTP'}), 400
        
        # Verify OTP
        if otp_data['otp'] == otp:
            # OTP is correct, generate reset token
            reset_token = secrets.token_urlsafe(32)
            otp_data['reset_token'] = reset_token
            otp_data['reset_expiry'] = datetime.now() + timedelta(minutes=30)
            
            return jsonify({
                'message': 'OTP verified successfully',
                'reset_token': reset_token,
                'type': otp_type
            })
        else:
            # Increment attempts
            otp_data['attempts'] += 1
            return jsonify({'error': 'Invalid OTP'}), 400
            
    except Exception as e:
        logger.error(f"OTP verification error: {e}")
        return jsonify({'error': 'Failed to verify OTP'}), 500

@app.route('/api/otp/reset-password', methods=['POST'])
def reset_password():
    """Reset password using reset token"""
    try:
        data = request.get_json()
        email = data.get('email')
        reset_token = data.get('reset_token')
        new_password = data.get('new_password')
        
        if not email or not reset_token or not new_password:
            return jsonify({'error': 'Email, reset token, and new password are required'}), 400
        
        # Validate password strength
        if len(new_password) < 6:
            return jsonify({'error': 'Password must be at least 6 characters long'}), 400
        
        # Check if reset token is valid
        if email not in otp_storage:
            return jsonify({'error': 'Invalid or expired reset token'}), 400
        
        otp_data = otp_storage[email]
        
        if 'reset_token' not in otp_data or 'reset_expiry' not in otp_data:
            return jsonify({'error': 'Invalid reset token'}), 400
        
        if otp_data['reset_token'] != reset_token:
            return jsonify({'error': 'Invalid reset token'}), 400
        
        if datetime.now() > otp_data['reset_expiry']:
            del otp_storage[email]
            return jsonify({'error': 'Reset token has expired'}), 400
        
        # Update password in user storage (modify this based on your user storage)
        # For now, we'll update the ukpUsers in localStorage
        try:
            # This is a simplified approach - in production, you'd update your database
            # For now, we'll just return success and let the frontend handle the password update
            return jsonify({
                'message': 'Password reset successful',
                'email': email
            })
        except Exception as e:
            logger.error(f"Failed to update password: {e}")
            return jsonify({'error': 'Failed to update password'}), 500
        
        # Clean up OTP data
        del otp_storage[email]
        
    except Exception as e:
        logger.error(f"Password reset error: {e}")
        return jsonify({'error': 'Failed to reset password'}), 500

@app.route('/api/otp/create-account', methods=['POST'])
def create_account():
    """Create new account with OTP verification"""
    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')
        name = data.get('name', '')
        
        if not all([email, password]):
            return jsonify({'error': 'Email and password are required'}), 400
        
        # Validate password strength
        if len(password) < 6:
            return jsonify({'error': 'Password must be at least 6 characters long'}), 400
        
        # Check if user already exists (you can modify this based on your user storage)
        # For now, we'll assume the frontend handles this check
        
        # Generate OTP for account creation
        otp = generate_otp()
        expiry_time = datetime.now() + timedelta(minutes=10)
        
        # Store OTP with account creation data
        otp_storage[email] = {
            'otp': otp,
            'expiry': expiry_time,
            'attempts': 0,
            'type': 'account_creation',
            'password': password,
            'name': name
        }
        
        # Send OTP via email
        if send_otp_email(email, otp, 'account_creation'):
            return jsonify({
                'message': 'OTP sent for account creation',
                'email': email
            })
        else:
            return jsonify({'error': 'Failed to send OTP email'}), 500
            
    except Exception as e:
        logger.error(f"Error creating account: {e}")
        return jsonify({'error': 'Failed to create account'}), 500

@app.route('/api/otp/verify-account', methods=['POST'])
def verify_account():
    """Verify OTP and create account"""
    try:
        data = request.get_json()
        email = data.get('email')
        otp = data.get('otp')
        
        if not all([email, otp]):
            return jsonify({'error': 'Email and OTP are required'}), 400
        
        # Verify OTP
        if email not in otp_storage:
            return jsonify({'error': 'No OTP found for this email'}), 400
        
        otp_data = otp_storage[email]
        
        # Check if OTP is for account creation
        if otp_data.get('type') != 'account_creation':
            return jsonify({'error': 'Invalid OTP type'}), 400
        
        # Check expiry
        if datetime.now() > otp_data['expiry']:
            del otp_storage[email]
            return jsonify({'error': 'OTP has expired'}), 400
        
        # Check attempts
        if otp_data['attempts'] >= 3:
            del otp_storage[email]
            return jsonify({'error': 'Too many attempts. Please request a new OTP'}), 400
        
        # Verify OTP
        if otp_data['otp'] != otp:
            otp_data['attempts'] += 1
            return jsonify({'error': 'Invalid OTP'}), 400
        
        # Account creation successful
        # In production, you would save to database here
        # For now, we'll return success and let frontend handle user storage
        
        account_data = {
            'email': email,
            'password': otp_data['password'],
            'name': otp_data.get('name', ''),
            'created_at': datetime.now().isoformat()
        }
        
        # Clear OTP data
        del otp_storage[email]
        
        return jsonify({
            'message': 'Account created successfully',
            'account': account_data
        })
        
    except Exception as e:
        logger.error(f"Error verifying account: {e}")
        return jsonify({'error': 'Failed to verify account'}), 500

@app.route('/api/otp/delete-account', methods=['POST'])
def delete_account_otp():
    """Send OTP for account deletion"""
    try:
        data = request.get_json()
        email = data.get('email')
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
        
        # Check if user exists (you can modify this based on your user storage)
        # For now, we'll assume the user exists if they provide an email
        
        # Generate OTP for account deletion
        otp = generate_otp()
        expiry_time = datetime.now() + timedelta(minutes=10)
        
        # Store OTP with deletion data
        otp_storage[email] = {
            'otp': otp,
            'expiry': expiry_time,
            'attempts': 0,
            'type': 'account_deletion'
        }
        
        # Send OTP via email
        if send_otp_email(email, otp, 'account_deletion'):
            return jsonify({
                'message': 'OTP sent for account deletion',
                'email': email
            })
        else:
            return jsonify({'error': 'Failed to send OTP email'}), 500
            
    except Exception as e:
        logger.error(f"Error sending deletion OTP: {e}")
        return jsonify({'error': 'Failed to send deletion OTP'}), 500

@app.route('/api/otp/verify-deletion', methods=['POST'])
def verify_deletion_otp():
    """Verify OTP and delete account"""
    try:
        data = request.get_json()
        email = data.get('email')
        otp = data.get('otp')
        
        if not all([email, otp]):
            return jsonify({'error': 'Email and OTP are required'}), 400
        
        # Verify OTP
        if email not in otp_storage:
            return jsonify({'error': 'No OTP found for this email'}), 400
        
        otp_data = otp_storage[email]
        
        # Check if OTP is for account deletion
        if otp_data.get('type') != 'account_deletion':
            return jsonify({'error': 'Invalid OTP type'}), 400
        
        # Check expiry
        if datetime.now() > otp_data['expiry']:
            del otp_storage[email]
            return jsonify({'error': 'OTP has expired'}), 400
        
        # Check attempts
        if otp_data['attempts'] >= 3:
            del otp_storage[email]
            return jsonify({'error': 'Too many attempts. Please request a new OTP'}), 400
        
        # Verify OTP
        if otp_data['otp'] != otp:
            otp_data['attempts'] += 1
            return jsonify({'error': 'Invalid OTP'}), 400
        
        # Account deletion successful
        # In production, you would delete from database here
        # For now, we'll return success and let frontend handle user removal
        
        # Clear OTP data
        del otp_storage[email]
        
        return jsonify({
            'message': 'Account deleted successfully',
            'email': email
        })
        
    except Exception as e:
        logger.error(f"Error verifying deletion OTP: {e}")
        return jsonify({'error': 'Failed to verify deletion OTP'}), 500

# Banner management endpoints
@app.route('/api/banners/upload', methods=['POST', 'OPTIONS'])
def upload_banner():
    """Upload banner for a user"""
    if request.method == 'OPTIONS':
        return jsonify({'success': True})
    
    try:
        data = request.get_json()
        user_email = data.get('user_email')
        file_data = data.get('file_data')
        file_type = data.get('file_type')
        
        if not user_email or not file_data or not file_type:
            return jsonify({'success': False, 'error': 'Missing required fields'}), 400
        
        result = save_banner_file(user_email, file_data, file_type)
        
        if result['success']:
            return jsonify(result)
        else:
            return jsonify(result), 400
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/banners/<user_email>', methods=['GET'])
def get_banner(user_email):
    """Get banner for a user"""
    try:
        result = load_banner_file(user_email)
        
        if result and result.get('success'):
            return jsonify(result)
        else:
            return jsonify({'success': False, 'error': 'Banner not found'}), 404
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/banners/<user_email>', methods=['DELETE', 'OPTIONS'])
def delete_banner(user_email):
    """Delete banner for a user"""
    if request.method == 'OPTIONS':
        return jsonify({'success': True})
    
    try:
        result = remove_banner_file(user_email)
        
        if result['success']:
            return jsonify(result)
        else:
            return jsonify(result), 400
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/banners/storage/info', methods=['GET'])
def get_banner_storage():
    """Get banner storage information"""
    try:
        result = get_banner_storage_info()
        
        if result['success']:
            return jsonify(result)
        else:
            return jsonify(result), 400
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    print("Starting Flask app...")
    load_sessions()
    # Disable debug mode and reloader
    app.run(host='0.0.0.0', port=5000, debug=False, use_reloader=False) 